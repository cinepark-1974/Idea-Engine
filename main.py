"""
Idea Engine v2.0
BLUE JEANS PICTURES · Creative Discovery & Triage Engine

v1.0: TRIAGE 트랙 (7-Stage 진단·판정 → LOCKED 시드)
v1.1: Creator Engine v2.5.2 정합 — 5개 신규 LOCKED 키 출력
      (locked_core_decisions / locked_music_rules / locked_visual_motifs
       / locked_ending_form / locked_creator_questions)
v2.0: HUNTER 트랙 추가 (5개 입구 아이디어 발굴 엔진)

[모드 분기]
HOME → [HUNTER 발굴 | TRIAGE 진단] 선택
HUNTER 출력 시드 → TRIAGE Stage 1 자동 전달

[설계 철학]
"카탈로그를 보여주는 게 아니라 작가 안에 잠재된 답을 끌어내기"

Writer Engine v3.1 디자인 시스템 적용
"""

import json
import re
import io
import os
from datetime import datetime
from typing import Dict, Any, Optional

import streamlit as st
import plotly.graph_objects as go
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import prompt as P
import market_lens_pack as MLP

# ─────────────────────────────────────
# Engine Info
# ─────────────────────────────────────
ENGINE_VERSION = "v2.0"
ENGINE_BUILD_DATE = "2026-05-19"
ENGINE_PATCH_LEVEL = "v2.0 (HUNTER 골격) + v1.4.1 패치 (Market Lens — KR·JP·ID 3개 시장 좌표 + Stage 6 UI 동적 라벨)"

ANTHROPIC_MODEL_SONNET = "claude-sonnet-4-6"
ANTHROPIC_MODEL_OPUS = "claude-opus-4-7"
MAX_TOKENS = 16000

# ─────────────────────────────────────
# Page Config
# ─────────────────────────────────────
st.set_page_config(
    page_title="BLUE JEANS · Idea Engine",
    page_icon="💡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ═══════════════════════════════════════════════════════════
# Session State + Mode-Switch Helpers
# (사이드바보다 먼저 정의되어야 함 — 사이드바 버튼이 호출함)
# ═══════════════════════════════════════════════════════════
def init_session_state():
    defaults = {
        # ── v2.0 모드 분기 ──
        "mode": "HOME",                # HOME | HUNTER | TRIAGE
        # ── HUNTER 트랙 (v2.0 신규) ──
        "hunter_entry": None,          # None | "0" | "1" | "2" | "3" | "4" | "5"
        "hunter_input": "",            # 입구 0 자유 텍스트 또는 입구별 입력
        "hunter_classified": None,     # 입구 0 자동분류 결과 (Sonnet)
        "hunter_stage_data": {},       # 입구별 진행 데이터 (질문 응답, 시드 후보 등)
        "hunter_output": None,         # 최종 LOCKED 시드 JSON (TRIAGE 입력으로 전달)
        # ── TRIAGE 트랙 (v1.0 호환) ──
        "current_stage": 1,
        "stage_1_input": None,
        "stage_2_logline": None,
        "stage_3_hook": None,
        "stage_4_format": None,
        "stage_5_reference": None,
        "stage_6_market": None,
        "stage_7_verdict": None,
        "selected_logline": None,
        "seed_loaded_from_hunter": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def reset_session():
    keys_to_clear = [
        "mode",
        "hunter_entry", "hunter_input", "hunter_classified",
        "hunter_stage_data", "hunter_output",
        "current_stage", "stage_1_input", "stage_2_logline",
        "stage_3_hook", "stage_4_format", "stage_5_reference",
        "stage_6_market", "stage_7_verdict", "selected_logline",
        "seed_loaded_from_hunter",
    ]
    for k in keys_to_clear:
        if k in st.session_state:
            del st.session_state[k]
    init_session_state()


def reset_triage_only():
    """TRIAGE 트랙만 리셋 (HUNTER 결과는 보존)."""
    keys = [
        "current_stage", "stage_1_input", "stage_2_logline",
        "stage_3_hook", "stage_4_format", "stage_5_reference",
        "stage_6_market", "stage_7_verdict", "selected_logline",
        "seed_loaded_from_hunter",
    ]
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]
    st.session_state["current_stage"] = 1
    st.session_state["seed_loaded_from_hunter"] = False


def reset_hunter_only():
    """HUNTER 트랙만 리셋 (TRIAGE 진행 상태는 보존)."""
    keys = [
        "hunter_entry", "hunter_input", "hunter_classified",
        "hunter_stage_data", "hunter_output",
    ]
    for k in keys:
        if k in st.session_state:
            del st.session_state[k]
    st.session_state["hunter_entry"] = None
    st.session_state["hunter_input"] = ""
    st.session_state["hunter_stage_data"] = {}


def transfer_hunter_seed_to_triage():
    """HUNTER 출력 시드를 TRIAGE Stage 1 입력 형식으로 변환 후 전달.

    HUNTER 시드 JSON 예상 스키마:
    {
        "title": str, "genre": str, "target_market": str,
        "format_pref": str, "raw_idea": str,
        "hunter_meta": {"entry": "1"~"5", "discovery_notes": [...]}
    }
    """
    seed = st.session_state.get("hunter_output")
    if not seed:
        return False

    st.session_state["stage_1_input"] = {
        "title": seed.get("title", "(제목 미정)"),
        "genre": seed.get("genre", "장르 미정"),
        "target_market": seed.get("target_market", "한국 + 글로벌"),
        "format": seed.get("format_pref", "미정 (Idea Engine이 추천)"),
        "raw_idea": seed.get("raw_idea", ""),
        "_hunter_meta": seed.get("hunter_meta", {}),
    }
    st.session_state["current_stage"] = 1
    st.session_state["seed_loaded_from_hunter"] = True
    return True


# ═══════════════════════════════════════════════════════════
# 진행 상태 JSON 저장·복원 (TRIAGE 중간 백업)
# 각 Stage 완료 후 다운로드, Stage 1 진입 시 업로드 복원
# ═══════════════════════════════════════════════════════════

PROGRESS_STAGE_KEYS = [
    "stage_1_input", "stage_2_logline", "stage_3_hook", "stage_4_format",
    "stage_5_reference", "stage_6_market", "stage_7_verdict", "selected_logline",
]
PROGRESS_SCHEMA = "triage_progress_v1"


def _detect_last_completed_stage() -> int:
    """마지막 완료 Stage 번호 (0=없음)."""
    pairs = [
        (1, "stage_1_input"), (2, "stage_2_logline"), (3, "stage_3_hook"),
        (4, "stage_4_format"), (5, "stage_5_reference"), (6, "stage_6_market"),
        (7, "stage_7_verdict"),
    ]
    last = 0
    for n, k in pairs:
        if st.session_state.get(k):
            last = n
        else:
            break
    return last


def build_progress_json(state: Dict[str, Any]) -> str:
    """현재 TRIAGE 진행 상태를 JSON 문자열로 직렬화."""
    last = _detect_last_completed_stage()
    title = ""
    s1 = state.get("stage_1_input")
    if isinstance(s1, dict):
        title = s1.get("title", "")
    progress = {
        "_idea_engine_progress": {
            "version": ENGINE_VERSION,
            "build_date": ENGINE_BUILD_DATE,
            "saved_at": datetime.now().isoformat(),
            "last_completed_stage": last,
            "project_title": title,
            "schema": PROGRESS_SCHEMA,
        },
    }
    for k in PROGRESS_STAGE_KEYS:
        progress[k] = state.get(k)
    # HUNTER 인계 이력 보존
    progress["_hunter_trace"] = {
        "seed_loaded_from_hunter": state.get("seed_loaded_from_hunter", False),
        "hunter_output": state.get("hunter_output"),
    }
    return json.dumps(progress, ensure_ascii=False, indent=2)


def load_progress_json(uploaded_dict: Dict[str, Any]) -> tuple:
    """업로드 JSON을 session_state에 복원. (success, message, last_stage)."""
    meta = uploaded_dict.get("_idea_engine_progress", {})
    if meta.get("schema") != PROGRESS_SCHEMA:
        return False, f"호환 안 됨 (필요: {PROGRESS_SCHEMA})", 0
    last = meta.get("last_completed_stage", 0)
    if not isinstance(last, int) or last < 1:
        return False, "유효하지 않은 last_completed_stage", 0
    for k in PROGRESS_STAGE_KEYS:
        if k in uploaded_dict:
            st.session_state[k] = uploaded_dict[k]
    htr = uploaded_dict.get("_hunter_trace", {})
    if isinstance(htr, dict):
        st.session_state["seed_loaded_from_hunter"] = htr.get("seed_loaded_from_hunter", False)
        if htr.get("hunter_output"):
            st.session_state["hunter_output"] = htr["hunter_output"]
    next_stage = min(last + 1, 7) if last < 7 else 7
    st.session_state["current_stage"] = next_stage
    st.session_state["mode"] = "TRIAGE"
    title = meta.get("project_title", "(제목 없음)")
    saved = meta.get("saved_at", "")
    return True, f"✓ '{title}' 복원 완료 — Stage {last}까지 완료. Stage {next_stage}부터 진행. (저장: {saved})", last


def render_progress_save_button(stage_num: int):
    """진행 상태 JSON 다운로드 버튼 (각 Stage 완료 페이지에 호출)."""
    state = dict(st.session_state)
    last = _detect_last_completed_stage()
    if last < 1:
        return
    title = ""
    s1 = state.get("stage_1_input")
    if isinstance(s1, dict):
        title = s1.get("title", "untitled").replace(" ", "_").replace("/", "_")[:40]
    json_str = build_progress_json(state)
    filename = f"IdeaProgress_{title}_stage{last}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    st.download_button(
        label=f"💾 진행 상태 백업 (Stage {last}까지) — JSON 다운로드",
        data=json_str.encode("utf-8"),
        file_name=filename,
        mime="application/json",
        key=f"progress_save_stage_{stage_num}",
        use_container_width=True,
        help="Stage 7 진단 전 백업 권장. 에러 발생 시 이 파일로 복원 가능.",
    )


def render_progress_load_widget():
    """진행 상태 JSON 업로드 위젯 (Stage 1 상단에 호출)."""
    with st.expander("📂 이전 진행 상태 JSON 복원하기", expanded=False):
        st.caption("이전에 백업한 IdeaProgress_*.json 파일을 업로드하면 그 단계로 복원됩니다.")
        uploaded = st.file_uploader(
            "JSON 파일 업로드",
            type=["json"],
            key="progress_load_uploader",
            label_visibility="collapsed",
        )
        if uploaded is not None:
            try:
                data = json.load(uploaded)
                if st.button("→ 이 JSON으로 복원", key="progress_load_btn", type="primary", use_container_width=True):
                    ok, msg, last = load_progress_json(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
            except json.JSONDecodeError as e:
                st.error(f"JSON 파싱 실패: {e}")
            except Exception as e:
                st.error(f"복원 실패: {type(e).__name__}: {e}")


init_session_state()

# ─────────────────────────────────────
# Sidebar Engine Info (Writer Engine 동일)
# ─────────────────────────────────────
with st.sidebar:
    st.markdown(f"""
    <div style="padding:12px;background:#F0F2FF;border-radius:8px;border-left:3px solid #191970;font-family:'Pretendard',sans-serif;">
        <div style="font-size:.72rem;color:#191970;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">ENGINE INFO</div>
        <div style="font-size:1.05rem;font-weight:700;color:#191970;">Idea Engine</div>
        <div style="font-size:1.25rem;font-weight:900;color:#FFCB05;background:#191970;padding:2px 8px;border-radius:4px;display:inline-block;margin-top:4px;">
            {ENGINE_VERSION}
        </div>
        <div style="font-size:.7rem;color:#666;margin-top:8px;">
            Build: {ENGINE_BUILD_DATE}<br>
            HUNTER 발굴 + TRIAGE 진단<br>
            <span style="color:#191970;font-weight:600;">+ v1.1 Creator v2.5.2 정합 5키</span><br>
            <span style="color:#191970;font-weight:600;">+ v1.2 Stanton 5원칙 · Hook&Punch 4키</span><br>
            <span style="color:#191970;font-weight:600;">+ v1.3 장르 · 시장 좌표 2키</span><br>
            <span style="color:#191970;font-weight:600;">+ v1.4.1 Market Lens (KR·JP·ID) + UI 동적</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

    # ── 모드 전환 패널 (v2.0 신규) ──
    current_mode = st.session_state.get("mode", "HOME")
    mode_label_map = {
        "HOME": "🏠 HOME",
        "HUNTER": "🎯 HUNTER (발굴)",
        "TRIAGE": "🔍 TRIAGE (진단)",
    }
    st.markdown(f"""
    <div style="padding:10px;background:#191970;border-radius:8px;font-family:'Pretendard',sans-serif;">
        <div style="font-size:.7rem;color:#FFCB05;font-weight:700;letter-spacing:.05em;margin-bottom:4px;">CURRENT MODE</div>
        <div style="font-size:.95rem;color:#fff;font-weight:700;">{mode_label_map[current_mode]}</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown("###### 🔀 모드 전환")

    if st.button("🏠 홈", key="nav_home", use_container_width=True):
        st.session_state["mode"] = "HOME"
        st.rerun()
    if st.button("🎯 HUNTER 트랙", key="nav_hunter", use_container_width=True):
        st.session_state["mode"] = "HUNTER"
        st.session_state["hunter_entry"] = None
        st.rerun()
    if st.button("🔍 TRIAGE 트랙", key="nav_triage", use_container_width=True):
        st.session_state["mode"] = "TRIAGE"
        st.rerun()

    # ── HUNTER → TRIAGE 시드 전송 (시드 준비된 경우만 노출) ──
    if st.session_state.get("hunter_output"):
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        st.markdown("""
        <div style="padding:8px 10px;background:#FFCB05;border-radius:6px;font-family:'Pretendard',sans-serif;font-size:.75rem;color:#191970;font-weight:700;">
            🔗 HUNTER 시드 준비됨
        </div>
        """, unsafe_allow_html=True)
        if st.button("→ TRIAGE로 전송", key="seed_to_triage", use_container_width=True, type="primary"):
            transfer_hunter_seed_to_triage()
            st.session_state["mode"] = "TRIAGE"
            st.rerun()

    # ── TRIAGE 진행 중 백업 (Stage 1 이상 데이터가 있을 때만) ──
    if _detect_last_completed_stage() >= 1:
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        last_stage_num = _detect_last_completed_stage()
        st.markdown(f"""
        <div style="padding:8px 10px;background:#E8F5E9;border-radius:6px;font-family:'Pretendard',sans-serif;font-size:.75rem;color:#1B5E20;font-weight:700;">
            💾 진행 중 (Stage {last_stage_num} 완료)
        </div>
        """, unsafe_allow_html=True)
        # 사이드바용 작은 버튼
        _state_for_sidebar = dict(st.session_state)
        _title_for_sidebar = ""
        _s1 = _state_for_sidebar.get("stage_1_input")
        if isinstance(_s1, dict):
            _title_for_sidebar = _s1.get("title", "untitled").replace(" ", "_").replace("/", "_")[:30]
        _sb_json = build_progress_json(_state_for_sidebar)
        _sb_filename = f"IdeaProgress_{_title_for_sidebar}_stage{last_stage_num}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        st.download_button(
            label="💾 백업 JSON 다운로드",
            data=_sb_json.encode("utf-8"),
            file_name=_sb_filename,
            mime="application/json",
            key="sidebar_progress_save",
            use_container_width=True,
            help="현재까지 진행한 모든 Stage 데이터를 JSON으로 저장",
        )

    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

    st.markdown("""
    <div style="padding:10px;background:#FFF8DC;border-radius:8px;font-family:'Pretendard',sans-serif;font-size:.78rem;">
        <div style="font-weight:700;color:#191970;margin-bottom:4px;">🤖 모델 정책</div>
        <div style="color:#444;">진단 ②~⑥: <b>Sonnet 4.6</b></div>
        <div style="color:#444;">최종 판정 ⑦: <b>Opus 4.7</b></div>
        <div style="color:#444;">HUNTER 발굴: <b>Sonnet 4.6</b></div>
    </div>
    """, unsafe_allow_html=True)

    st.caption("버전이 최신인지 확인하세요.")

# ─────────────────────────────────────
# Custom CSS (Writer Engine v3.1 동일 톤)
# ─────────────────────────────────────
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #191970; --y: #FFCB05; --bg: #F7F7F5;
    --card: #FFFFFF; --card-border: #E2E2E0; --t: #1A1A2E;
    --g: #2EC484; --r: #D33F49; --o: #E8B800;
    --dim: #8E8E99; --light-bg: #EEEEF6;
    --serif: 'Paperlogy', 'Noto Serif KR', 'Georgia', serif;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--body); color: var(--t); -webkit-font-smoothing: antialiased;
}
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important; color: var(--t) !important;
}
.stMarkdown, .stText, .stCode { color: var(--t) !important; }
h1,h2,h3,h4,h5,h6 { color: var(--navy) !important; font-family: var(--heading) !important; }
p, span, label, div, li { color: inherit; }

.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-size: 0.92rem !important;
    padding: 0.65rem 0.85rem !important;
}
.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(25,25,112,0.08) !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important; font-size: 0.85rem !important;
}
.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important; color: var(--t) !important;
    border-color: var(--card-border) !important; border-radius: 8px !important;
}
[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important; color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }
.stTextInput label, .stTextArea label, .stSelectbox label, .stRadio label, .stFileUploader label {
    color: var(--t) !important; font-weight: 600 !important;
    font-size: 0.82rem !important; margin-bottom: 0.3rem !important;
}

.stButton > button {
    color: var(--t) !important; border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-weight: 700 !important;
    font-size: 0.88rem !important; padding: 0.55rem 1.2rem !important;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(25,25,112,0.08) !important;
}
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important; color: var(--navy) !important;
    border-color: var(--y) !important; font-weight: 800 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.3) !important;
}
.stDownloadButton > button {
    color: var(--navy) !important; border: 1.5px solid var(--y) !important;
    background-color: var(--y) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-weight: 800 !important;
    font-size: 0.88rem !important; padding: 0.55rem 1.2rem !important;
}
.stExpander, details, details summary {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1px solid var(--card-border) !important; border-radius: 8px !important;
}
details[open] > div { background-color: var(--card) !important; }
.stExpander summary, .stExpander summary span { color: var(--t) !important; }
.stAlert { color: var(--t) !important; border-radius: 8px !important; }
[data-testid="stVerticalBlock"], [data-testid="stHorizontalBlock"],
[data-testid="stColumn"] { background-color: transparent !important; }

.header {
    font-size: 0.85rem; font-weight: 700; color: var(--navy);
    letter-spacing: 0.15em; font-family: var(--heading);
}
.brand-title {
    font-size: 2.6rem; font-weight: 900; color: var(--navy);
    font-family: var(--display); letter-spacing: -0.02em;
    position: relative; display: inline-block;
}
.brand-title::after {
    content: ''; position: absolute; bottom: 2px; left: 0;
    width: 100%; height: 4px; background: var(--y); border-radius: 2px;
}
.sub {
    font-size: 0.7rem; color: var(--dim); letter-spacing: 0.15em;
    margin-top: 0.5rem; margin-bottom: 1.5rem;
}
.callout {
    background: var(--light-bg); border-left: 4px solid var(--navy);
    padding: 0.9rem 1.1rem; margin: 0.5rem 0;
    border-radius: 0 8px 8px 0; font-size: 0.88rem; color: var(--t);
}
.cl {
    color: var(--navy); font-weight: 700; font-size: 0.72rem;
    letter-spacing: 0.03em; margin-bottom: 0.3rem; text-transform: uppercase;
}
.section-header {
    background: var(--y); color: var(--navy);
    padding: 0.6rem 1rem; border-radius: 6px;
    font-weight: 800; font-size: 1rem; font-family: var(--heading);
    margin: 1.5rem 0 0.8rem 0;
    display: flex; justify-content: space-between; align-items: center;
}
.section-header .en {
    font-family: var(--display); font-size: 0.75rem;
    font-weight: 700; letter-spacing: 0.05em; opacity: 0.7;
}
.small-meta {
    font-size: 0.78rem; color: var(--dim);
    margin-top: -0.2rem; margin-bottom: 0.5rem;
}
.beat-tag {
    background: var(--navy); color: var(--y);
    display: inline-block; padding: 0.2rem 0.7rem;
    border-radius: 4px; font-size: 0.78rem; font-weight: 800;
    letter-spacing: 0.04em; margin-bottom: 0.4rem;
}
.act-tag {
    background: var(--navy); color: #fff;
    display: inline-block; padding: 0.25rem 0.8rem;
    border-radius: 4px; font-size: 0.82rem; font-weight: 800;
    letter-spacing: 0.06em;
}

/* ── Idea Engine 전용 컴포넌트 ── */
.stepper {
    display: flex; gap: 4px; margin: 1.5rem 0;
    padding: 0.8rem; background: var(--card);
    border-radius: 10px; border: 1px solid var(--card-border);
    overflow-x: auto;
}
.step {
    flex: 1; text-align: center; padding: 0.5rem 0.4rem;
    border-radius: 6px; font-size: 0.75rem;
    font-family: var(--heading); font-weight: 700;
    color: var(--dim); border: 1.5px solid transparent;
    min-width: 90px;
}
.step.active {
    background: var(--y); color: var(--navy);
    border-color: var(--y);
}
.step.done {
    background: var(--light-bg); color: var(--navy);
    border-color: var(--navy);
}
.step .num {
    font-size: 1rem; font-weight: 900; display: block;
    font-family: var(--display);
}

.score-card {
    background: var(--card); border: 1px solid var(--card-border);
    border-radius: 8px; padding: 0.9rem 1.1rem; margin: 0.5rem 0;
    border-left-width: 5px;
}
.score-card.pass { border-left-color: var(--g); }
.score-card.warn { border-left-color: var(--o); }
.score-card.fail { border-left-color: var(--r); }
.score-card .axis-name {
    font-weight: 800; font-family: var(--heading);
    color: var(--navy); font-size: 0.95rem;
}
.score-card .axis-score {
    font-family: var(--display); font-weight: 900;
    font-size: 1.4rem; margin-left: 0.5rem;
}
.score-card .axis-comment {
    color: var(--dim); font-size: 0.85rem; margin-top: 0.2rem;
}

.metric-tile {
    background: var(--card); border: 1px solid var(--card-border);
    border-radius: 10px; padding: 1rem; text-align: center;
}
.metric-tile .num {
    font-family: var(--display); font-size: 2.2rem; font-weight: 900;
    color: var(--navy); line-height: 1;
}
.metric-tile .label {
    font-size: 0.75rem; color: var(--dim); margin-top: 0.4rem;
    text-transform: uppercase; letter-spacing: 0.08em;
    font-family: var(--heading); font-weight: 700;
}

.verdict-box {
    padding: 2.5rem 2rem; border-radius: 12px; margin: 1.5rem 0;
    text-align: center; font-family: var(--display);
    border: 3px solid;
}
.verdict-box.go { background: linear-gradient(135deg,#E8F5E9,#C8E6C9); border-color: var(--g); }
.verdict-box.cond { background: linear-gradient(135deg,#FFF8E1,#FFECB3); border-color: var(--o); }
.verdict-box.nogo { background: linear-gradient(135deg,#FFEBEE,#FFCDD2); border-color: var(--r); }
.verdict-box .verdict-label {
    font-size: 3rem; font-weight: 900; letter-spacing: 0.12em;
    margin: 0; line-height: 1;
}

.locked-card {
    background: var(--navy); color: white;
    padding: 1.5rem 1.7rem; border-radius: 10px; margin: 1rem 0;
    border-left: 6px solid var(--y);
}
.locked-card .field-label {
    color: var(--y); font-size: 0.72rem;
    text-transform: uppercase; letter-spacing: 0.1em;
    margin-bottom: 0.3rem; font-family: var(--heading); font-weight: 700;
}
.locked-card .field-value {
    color: white; font-size: 0.95rem; margin-bottom: 1rem;
    line-height: 1.5;
}
.locked-card .field-value:last-child { margin-bottom: 0; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────
# Anthropic Client + JSON Parser
# ─────────────────────────────────────
def get_anthropic_client() -> Optional[Anthropic]:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", os.getenv("ANTHROPIC_API_KEY"))
    return Anthropic(api_key=api_key) if api_key else None


def safe_json_loads(text: str) -> Dict[str, Any]:
    """4단계 JSON 복구 시스템"""
    text = text.strip()
    text = re.sub(r"^```json\s*", "", text)
    text = re.sub(r"^```\s*", "", text)
    text = re.sub(r"\s*```$", "", text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    try:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1:
            return json.loads(text[start:end + 1])
    except json.JSONDecodeError:
        pass

    try:
        candidate = text[text.find("{"):text.rfind("}") + 1]
        for _ in range(30):
            try:
                return json.loads(candidate)
            except json.JSONDecodeError as e:
                pos = e.pos
                if 0 < pos < len(candidate):
                    candidate = candidate[:pos] + " " + candidate[pos + 1:]
                else:
                    break
    except Exception:
        pass

    try:
        cleaned = re.sub(r'(?<="):\s*"([^"]*)"\s*([,\}])',
                         lambda m: f': "{m.group(1).replace(chr(34), chr(39))}" {m.group(2)}',
                         text)
        return json.loads(cleaned)
    except Exception:
        return {"_parse_error": True, "_raw": text}


def call_claude(client: Anthropic, prompt_text: str, model: str = ANTHROPIC_MODEL_SONNET) -> Dict[str, Any]:
    response = client.messages.create(
        model=model,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt_text}],
    )
    raw = response.content[0].text
    return safe_json_loads(raw)


# ─────────────────────────────────────
# DOCX Export
# ─────────────────────────────────────
def add_section_header_docx(doc, kr: str, en: str):
    p = doc.add_paragraph()
    run_kr = p.add_run(kr + " ")
    run_kr.font.size = Pt(14)
    run_kr.font.bold = True
    run_kr.font.color.rgb = RGBColor(0x19, 0x19, 0x70)
    run_en = p.add_run(en)
    run_en.font.size = Pt(11)
    run_en.font.italic = True
    run_en.font.color.rgb = RGBColor(0x6B, 0x6B, 0x7B)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFCB05')
    pPr.append(shd)


def add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def build_diagnostic_docx(state: Dict[str, Any]) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("IDEA DIAGNOSTIC REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x19, 0x19, 0x70)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("아이디어 진단 보고서")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x7B)
    
    doc.add_paragraph()
    inp = state["stage_1_input"]
    
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(inp["title"])
    run.font.size = Pt(20)
    run.font.bold = True
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"{inp['genre']} · {inp['format']} · {inp['target_market']}")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x7B)

    doc.add_paragraph()
    doc.add_paragraph()
    
    today = doc.add_paragraph()
    today.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = today.add_run(f"발행일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x7B)
    
    bjp = doc.add_paragraph()
    bjp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bjp.add_run(f"BLUE JEANS PICTURES · Idea Engine {ENGINE_VERSION}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x19, 0x19, 0x70)
    run.font.bold = True

    doc.add_page_break()

    add_section_header_docx(doc, "1. 원본 아이디어", "ORIGINAL IDEA")
    add_para(doc, inp["raw_idea"])
    doc.add_paragraph()

    if state["stage_2_logline"]:
        add_section_header_docx(doc, "2. 로그라인 정제", "LOGLINE REFINEMENT")
        ll = state["stage_2_logline"]
        for v in ll.get("logline_variants", []):
            add_para(doc, f"[{v['variant']}안 - {v['label']}]", bold=True)
            add_para(doc, v["logline"])
            add_para(doc, f"  강점: {v['strength']}", italic=True, size=10)
            add_para(doc, f"  약점: {v['weakness']}", italic=True, size=10)
            doc.add_paragraph()
        add_para(doc, f"▶ 추천: {ll.get('recommended', '')}안", bold=True)
        add_para(doc, ll.get("recommendation_reason", ""))
        doc.add_paragraph()

    if state["stage_3_hook"]:
        add_section_header_docx(doc, "3. 후크 진단 (Gate 0)", "HOOK DIAGNOSTIC")
        hk = state["stage_3_hook"]
        scores = hk.get("scores", {})
        
        tbl = doc.add_table(rows=6, cols=3)
        tbl.style = "Light Grid Accent 1"
        h = tbl.rows[0].cells
        h[0].text = "축"; h[1].text = "점수"; h[2].text = "코멘트"
        
        axis_kr = {
            "specificity": "구체성", "conflict_visibility": "갈등 가시성",
            "genre_clarity": "장르 명확성", "stakes": "판돈", "originality": "독창성"
        }
        for i, (k, kr) in enumerate(axis_kr.items(), 1):
            sc = scores.get(k, {})
            r = tbl.rows[i].cells
            r[0].text = kr
            r[1].text = f"{sc.get('score', 0)}/10"
            r[2].text = sc.get("comment", "")
        
        doc.add_paragraph()
        add_para(doc, f"총점: {hk.get('total_score', 0)}/50 — {hk.get('gate_status', '')}", bold=True, size=13)
        doc.add_paragraph()
        
        add_para(doc, "강점", bold=True)
        for s in hk.get("key_strengths", []):
            add_para(doc, f"  • {s}")
        doc.add_paragraph()
        add_para(doc, "약점", bold=True)
        for w in hk.get("key_weaknesses", []):
            add_para(doc, f"  • {w}")
        doc.add_paragraph()
        add_para(doc, "보강 제안", bold=True)
        for s in hk.get("improvement_suggestions", []):
            add_para(doc, f"  • {s}")
        doc.add_paragraph()

    if state["stage_4_format"]:
        add_section_header_docx(doc, "4. 포맷 추천", "FORMAT RECOMMENDATION")
        fm = state["stage_4_format"]
        fs = fm.get("format_scores", {})
        
        tbl = doc.add_table(rows=6, cols=3)
        tbl.style = "Light Grid Accent 1"
        h = tbl.rows[0].cells
        h[0].text = "포맷"; h[1].text = "점수"; h[2].text = "근거"
        
        format_kr = {
            "feature_film": "장편 영화", "ott_series": "OTT 시리즈",
            "mini_series": "미니시리즈", "short_form": "숏폼 드라마", "web_novel": "웹소설"
        }
        for i, (k, kr) in enumerate(format_kr.items(), 1):
            f = fs.get(k, {})
            r = tbl.rows[i].cells
            r[0].text = kr
            r[1].text = f"{f.get('score', 0)}/10"
            r[2].text = f.get("reason", "")
        
        doc.add_paragraph()
        primary = fm.get("primary_format_detail", {})
        add_para(doc, f"▶ 1순위: {primary.get('format_name', '')}", bold=True, size=13)
        if primary.get("episode_count"):
            add_para(doc, f"  회차: {primary['episode_count']}")
        if primary.get("runtime_per_episode"):
            add_para(doc, f"  회당 분량: {primary['runtime_per_episode']}")
        doc.add_paragraph()
        add_para(doc, "IP 빌딩 전략", bold=True)
        add_para(doc, fm.get("ip_building_strategy", ""))
        doc.add_paragraph()

    if state["stage_5_reference"]:
        add_section_header_docx(doc, "5. 레퍼런스 매핑", "REFERENCE MAPPING")
        rf = state["stage_5_reference"]
        for ref in rf.get("references", []):
            add_para(doc, f"《{ref['title']}》 ({ref.get('year', '')}, {ref.get('country', '')}) - {ref.get('format', '')}", bold=True)
            add_para(doc, f"  유사 차원: {ref.get('similarity_axis', '')}", italic=True, size=10)
            add_para(doc, f"  공통점: {ref.get('common_points', '')}")
            add_para(doc, f"  차별점: {ref.get('differentiation', '')}")
            doc.add_paragraph()
        warn = rf.get("lethal_similarity_warning", {})
        if warn.get("exists"):
            add_para(doc, "⚠ 치명적 유사작 경고", bold=True)
            add_para(doc, warn.get("details", ""))
        else:
            add_para(doc, "✓ 치명적 유사작 없음 - 안전", bold=True)
        doc.add_paragraph()
        add_para(doc, "차별화 요약", bold=True)
        add_para(doc, rf.get("differentiation_summary", ""))
        doc.add_paragraph()
        add_para(doc, "투자자 미팅 답변용", bold=True)
        add_para(doc, rf.get("investor_pitch_answer", ""))
        doc.add_paragraph()

    if state["stage_6_market"]:
        add_section_header_docx(doc, "6. 시장성 진단", "MARKET DIAGNOSTIC")
        mk = state["stage_6_market"]
        
        dom = mk.get("domestic_market", {})
        add_para(doc, f"한국 시장 ({'★' * dom.get('stars', 0)})", bold=True, size=13)
        ta = dom.get("target_audience", {})
        add_para(doc, f"  타겟: {ta.get('gender', '')} {ta.get('age_range', '')} - {ta.get('psychographic', '')}")
        add_para(doc, f"  예산: {dom.get('budget_estimate', '')}")
        add_para(doc, f"  유통: {', '.join(dom.get('distribution', []))}")
        add_para(doc, f"  IP 확장: {', '.join(dom.get('ip_extension_potential', []))}")
        doc.add_paragraph()
        
        glb = mk.get("global_market", {})
        add_para(doc, f"글로벌 시장 ({'★' * glb.get('stars', 0)})", bold=True, size=13)
        add_para(doc, f"  1차 타겟: {glb.get('primary_target_country', '')}")
        add_para(doc, f"  어필: {glb.get('global_appeal_strength', '')}")
        add_para(doc, f"  진입경로: {', '.join(glb.get('entry_path', []))}")
        add_para(doc, f"  약점: {glb.get('weakness', '')}")
        doc.add_paragraph()
        
        ott = mk.get("ott_market", {})
        add_para(doc, f"OTT 시장 ({'★' * ott.get('stars', 0)})", bold=True, size=13)
        fc = ott.get("first_choice_platform", {})
        sc = ott.get("second_choice_platform", {})
        add_para(doc, f"  1순위: {fc.get('name', '')} - {fc.get('reason', '')}")
        add_para(doc, f"  2순위: {sc.get('name', '')} - {sc.get('reason', '')}")
        add_para(doc, f"  최적 회차: {ott.get('optimal_episode_count', '')}")
        add_para(doc, f"  경쟁: {ott.get('competition_analysis', '')}")
        doc.add_paragraph()
        
        timing = mk.get("timing_fit", {})
        add_para(doc, f"시기적 적합성 ({'★' * timing.get('score', 0)})", bold=True)
        add_para(doc, timing.get("reason", ""))
        doc.add_paragraph()
        
        add_para(doc, "위험 신호", bold=True)
        for r in mk.get("risk_signals", []):
            add_para(doc, f"  • {r}")
        doc.add_paragraph()

    if state["stage_7_verdict"]:
        add_section_header_docx(doc, "7. 최종 판정", "FINAL VERDICT")
        vd = state["stage_7_verdict"]
        
        verdict = vd.get("final_verdict", "")
        vp = doc.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = vp.add_run(f"  {verdict}  ")
        run.font.size = Pt(28)
        run.font.bold = True
        if verdict == "GO":
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        elif verdict == "CONDITIONAL":
            run.font.color.rgb = RGBColor(0xF9, 0xA8, 0x25)
        else:
            run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        doc.add_paragraph()
        
        add_para(doc, "판정 사유", bold=True)
        add_para(doc, vd.get("verdict_reasoning", ""))
        doc.add_paragraph()
        
        if vd.get("conditional_requirements"):
            add_para(doc, "충족 조건 (CONDITIONAL)", bold=True)
            for c in vd["conditional_requirements"]:
                add_para(doc, f"  • {c}")
            doc.add_paragraph()
        
        if vd.get("nogo_alternative"):
            add_para(doc, "대안 제시 (NOGO)", bold=True)
            add_para(doc, vd["nogo_alternative"])
            doc.add_paragraph()
        
        add_para(doc, "확정된 핵심 결정", bold=True)
        for k in vd.get("key_decisions_made", []):
            add_para(doc, f"  • {k}")
        doc.add_paragraph()
        
        add_para(doc, "Creator Engine에서 결정해야 할 질문", bold=True)
        for q in vd.get("pending_decisions_for_creator", []):
            add_para(doc, f"  • {q}")
        doc.add_paragraph()
        
        add_section_header_docx(doc, "임원 요약", "EXECUTIVE SUMMARY")
        add_para(doc, vd.get("executive_summary", ""))
        doc.add_paragraph()
        
        add_section_header_docx(doc, "LOCKED 시드 패키지", "LOCKED SEED PACKAGE")
        add_para(doc, "Creator Engine 입력용 데이터 ─ 이 항목들은 Creator Engine에서 변경하지 않는다.", italic=True, size=10)
        doc.add_paragraph()
        
        seed = vd.get("locked_seed_package", {})
        add_para(doc, f"Project ID: {seed.get('project_id', '')}", bold=True, size=11)
        add_para(doc, f"제목 (KR): {seed.get('title_kr', '')}")
        add_para(doc, f"제목 (EN): {seed.get('title_en', '')}")
        doc.add_paragraph()
        
        add_para(doc, "LOCKED LOGLINE", bold=True)
        add_para(doc, seed.get("locked_logline", ""))
        doc.add_paragraph()
        
        add_para(doc, "LOCKED GENRE", bold=True)
        gn = seed.get("locked_genre", {})
        add_para(doc, f"  Primary: {gn.get('primary', '')}")
        add_para(doc, f"  Secondary: {gn.get('secondary', '')}")
        if gn.get("tertiary"):
            add_para(doc, f"  Tertiary: {gn['tertiary']}")
        doc.add_paragraph()
        
        add_para(doc, "LOCKED FORMAT", bold=True)
        ft = seed.get("locked_format", {})
        add_para(doc, f"  Primary: {ft.get('primary', '')}")
        if ft.get("episode_count"):
            add_para(doc, f"  Episodes: {ft['episode_count']}")
        if ft.get("runtime"):
            add_para(doc, f"  Runtime: {ft['runtime']}")
        add_para(doc, f"  IP Strategy: {ft.get('ip_strategy', '')}")
        doc.add_paragraph()
        
        add_para(doc, "LOCKED TARGET", bold=True)
        tg = seed.get("locked_target", {})
        add_para(doc, f"  Domestic: {tg.get('domestic', '')}")
        add_para(doc, f"  Global: {tg.get('global', '')}")
        doc.add_paragraph()
        
        add_para(doc, "LOCKED THEME", bold=True)
        th = seed.get("locked_theme", {})
        add_para(doc, f"  Surface: {th.get('surface', '')}")
        add_para(doc, f"  Deep: {th.get('deep', '')}")
        doc.add_paragraph()
        
        add_para(doc, "LOCKED REFERENCES", bold=True)
        for r in seed.get("locked_references", []):
            add_para(doc, f"  • {r}")
        doc.add_paragraph()
        
        add_para(doc, f"Hook Score: {seed.get('locked_hook_score', '')}/50", bold=True)
        ms = seed.get("locked_market_stars", {})
        add_para(doc, f"Market Stars: 한국 {'★' * ms.get('domestic', 0)} / 글로벌 {'★' * ms.get('global', 0)} / OTT {'★' * ms.get('ott', 0)}")
        add_para(doc, f"Distribution Priority: {seed.get('locked_distribution_priority', '')}")
        doc.add_paragraph()
        
        add_para(doc, "Creator Engine 진행 시 다뤄야 할 위험", bold=True)
        for risk in seed.get("locked_risks_to_address", []):
            add_para(doc, f"  • {risk}")

        # ════════════════════════════════════════════════════════
        # v1.1 — Creator Engine v2.5.2 정합 5개 신규 LOCKED 영역
        # ════════════════════════════════════════════════════════
        doc.add_paragraph()
        add_section_header_docx(
            doc, "v1.1 신규 LOCKED 영역",
            "CREATOR ENGINE v2.5.2 ABSORPTION KEYS"
        )
        add_para(
            doc,
            "Creator Engine v2.5.2가 작품 본질로 절대 보존하는 5개 영역. "
            "「오랜만에」 검증에서 발견된 핵심 모티프 휘발(61%)을 차단하기 위해 도입.",
            italic=True, size=10,
        )
        doc.add_paragraph()

        # ① locked_core_decisions
        core_decisions = seed.get("locked_core_decisions", []) or []
        add_para(doc, f"① 확정된 핵심 결정 (locked_core_decisions) — {len(core_decisions)}건", bold=True)
        if not core_decisions:
            add_para(doc, "  (이 작품에는 별도로 LOCK된 핵심 결정이 없음 — 빈 배열)", italic=True, size=10)
        else:
            for d in core_decisions:
                if isinstance(d, dict):
                    cat = d.get("category", "")
                    rule = d.get("rule", "") or d.get("decision", "")
                    rationale = d.get("rationale", "")
                    if cat and rule:
                        add_para(doc, f"  • [{cat}] {rule}")
                    elif rule:
                        add_para(doc, f"  • {rule}")
                    if rationale:
                        add_para(doc, f"      근거: {rationale}", italic=True, size=10)
                elif isinstance(d, str):
                    add_para(doc, f"  • {d}")
        doc.add_paragraph()

        # ② locked_music_rules
        music_rules = seed.get("locked_music_rules", {}) or {}
        add_para(
            doc,
            f"② 음악 사용 규약 (locked_music_rules) — {'있음' if music_rules else '없음 (빈 객체)'}",
            bold=True,
        )
        if not music_rules:
            add_para(doc, "  (이 작품에는 음악 사용 규약이 없음 — 액션·스릴러·호러에서 자주 발생)", italic=True, size=10)
        elif isinstance(music_rules, dict):
            for k, v in music_rules.items():
                if isinstance(v, list):
                    add_para(doc, f"  • [{k}]")
                    for item in v:
                        add_para(doc, f"      - {item}")
                else:
                    add_para(doc, f"  • [{k}] {v}")
        elif isinstance(music_rules, list):
            for r in music_rules:
                add_para(doc, f"  • {r}")
        else:
            add_para(doc, f"  • {music_rules}")
        doc.add_paragraph()

        # ③ locked_visual_motifs
        visual_motifs = seed.get("locked_visual_motifs", []) or []
        add_para(doc, f"③ 시각 모티프 (locked_visual_motifs) — {len(visual_motifs)}건", bold=True)
        if not visual_motifs:
            add_para(doc, "  (이 작품에는 LOCK된 시각 모티프가 없음 — 빈 배열)", italic=True, size=10)
        else:
            for m in visual_motifs:
                if isinstance(m, dict):
                    motif = m.get("motif", "") or m.get("name", "")
                    function = m.get("function", "") or m.get("role", "")
                    if motif and function:
                        add_para(doc, f"  • {motif} → {function}")
                    elif motif:
                        add_para(doc, f"  • {motif}")
                elif isinstance(m, str):
                    add_para(doc, f"  • {m}")
        doc.add_paragraph()

        # ④ locked_ending_form
        ending_form = seed.get("locked_ending_form", {}) or {}
        add_para(
            doc,
            f"④ 결말 형식 (locked_ending_form) — {'LOCK됨' if ending_form else '미확정 (빈 객체)'}",
            bold=True,
        )
        if not ending_form:
            add_para(doc, "  (결말 형식이 LOCK되지 않음 — Creator Engine이 결정)", italic=True, size=10)
        elif isinstance(ending_form, dict):
            if ending_form.get("type"):
                add_para(doc, f"  • 결말 유형: {ending_form['type']}")
            if ending_form.get("emotional_resolution"):
                add_para(doc, f"  • 정서적 해소: {ending_form['emotional_resolution']}")
            if ending_form.get("final_image"):
                add_para(doc, f"  • 마지막 이미지: {ending_form['final_image']}")
            if ending_form.get("forbidden"):
                add_para(doc, f"  • 금지 패턴: {ending_form['forbidden']}")
        else:
            add_para(doc, f"  • {ending_form}")
        doc.add_paragraph()

        # ⑤ locked_creator_questions
        creator_questions = seed.get("locked_creator_questions", []) or []
        add_para(
            doc,
            f"⑤ Creator Engine 결정 의제 (locked_creator_questions) — {len(creator_questions)}건",
            bold=True,
        )
        if not creator_questions:
            add_para(doc, "  (Creator Engine이 답할 의제가 없음 — 빈 배열)", italic=True, size=10)
        else:
            for q in creator_questions:
                if isinstance(q, dict):
                    question = q.get("question", "")
                    options = q.get("options", []) or []
                    importance = q.get("importance", "")
                    line = f"  • {question}"
                    if importance:
                        line += f"  [중요도: {importance.upper()}]"
                    add_para(doc, line)
                    if options:
                        add_para(doc, f"      후보: {' / '.join(str(o) for o in options)}", italic=True, size=10)
                elif isinstance(q, str):
                    add_para(doc, f"  • {q}")

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"© 2026 BLUE JEANS PICTURES · Idea Engine {ENGINE_VERSION}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x7B)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_seed_json(state: Dict[str, Any]) -> str:
    if not state.get("stage_7_verdict"):
        return "{}"
    seed = state["stage_7_verdict"].get("locked_seed_package", {})

    # ─── v1.1: Creator Engine v2.5.2 정합 5개 신규 키 빈 값 보장 ───
    # 작품 특성상 해당 없는 영역도 키 자체는 명시 출력 (Creator Engine이
    # "Idea Engine이 의식적으로 비웠다"는 신호로 해석)
    seed.setdefault("locked_core_decisions", [])
    seed.setdefault("locked_music_rules", {})
    seed.setdefault("locked_visual_motifs", [])
    seed.setdefault("locked_ending_form", {})
    seed.setdefault("locked_creator_questions", [])

    # ─── v1.2: Stanton 5원칙 + Hook & Punch 발굴 4개 신규 키 ───
    seed.setdefault("locked_empathy_anchor", {})
    seed.setdefault("locked_hook_signature", {})
    seed.setdefault("locked_punch_scene", {})
    seed.setdefault("locked_ending_promise", {})

    # ─── v1.3: 장르 + 시장 좌표 2개 신규 키 ───
    seed.setdefault("locked_genre_primary", {})
    seed.setdefault("locked_market_position", {})

    creator_input = {
        "_idea_engine_meta": {
            "version": ENGINE_VERSION,
            "patch": "v1.4.1 (Market Lens KR·JP·ID + UI 동적 라벨) on v1.3 (장르+시장좌표 2키) on v1.2 (Stanton+Hook&Punch 4키) on v1.1 (Creator v2.5.2 정합 5키)",
            "generated_at": datetime.now().isoformat(),
            "project_id": seed.get("project_id", ""),
            "verdict": state["stage_7_verdict"].get("final_verdict", ""),
            "hook_score": seed.get("locked_hook_score", 0),
        },
        "title": seed.get("title_kr", ""),
        "raw_idea": seed.get("locked_logline", ""),
        "genre": seed.get("locked_genre", {}).get("primary", ""),
        "target_market": seed.get("locked_target", {}).get("domestic", ""),
        "format": seed.get("locked_format", {}).get("primary", ""),
        "locked_seed": seed,
        "executive_summary": state["stage_7_verdict"].get("executive_summary", ""),
        "pending_decisions": state["stage_7_verdict"].get("pending_decisions_for_creator", []),
    }
    return json.dumps(creator_input, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════════
# HEADER (Writer Engine 양식 동일)
# ═══════════════════════════════════════════════════════════
st.markdown(
    '<div style="text-align:center;padding:1rem 0 0 0">'
    '<div class="header">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>'
    '<div class="brand-title">IDEA ENGINE</div>'
    '<div class="sub">Y O U N G &nbsp; · &nbsp; V I N T A G E &nbsp; · &nbsp; F R E E &nbsp; · &nbsp; I N N O V A T I V E</div>'
    '</div>',
    unsafe_allow_html=True,
)


# ─────────────────────────────────────
# Stepper
# ─────────────────────────────────────
def render_stepper(current: int):
    stages_meta = [
        (1, "아이디어"), (2, "로그라인"), (3, "Hook"),
        (4, "Format"), (5, "Reference"), (6, "Market"), (7, "최종 판정")
    ]
    completed_keys = {
        1: "stage_1_input", 2: "stage_2_logline", 3: "stage_3_hook",
        4: "stage_4_format", 5: "stage_5_reference", 6: "stage_6_market",
        7: "stage_7_verdict"
    }
    
    html = '<div class="stepper">'
    for num, name in stages_meta:
        completed = bool(st.session_state.get(completed_keys[num]))
        if num == current:
            cls = "step active"
            label_num = f"<span class='num'>{num}</span>"
        elif completed:
            cls = "step done"
            label_num = "<span class='num'>✓</span>"
        else:
            cls = "step"
            label_num = f"<span class='num'>{num}</span>"
        html += f'<div class="{cls}">{label_num}{name}</div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


def section_header(kr: str, en: str):
    st.markdown(
        f'<div class="section-header">{kr} <span class="en">{en}</span></div>',
        unsafe_allow_html=True,
    )


def small_meta(text: str):
    st.markdown(f'<div class="small-meta">{text}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────
# STAGE PAGES
# ─────────────────────────────────────
def page_stage_1():
    section_header("📥 STEP 1 · 아이디어 입력", "FROM RAW IDEA")
    small_meta("모호한 아이디어 한 줄부터 한 단락까지 자유롭게 입력하세요. Idea Engine이 정제하여 Creator Engine이 받아먹을 수 있는 LOCKED 시드 패키지로 변환합니다.")

    # ── 이전 진행 상태 JSON 복원 위젯 ──
    render_progress_load_widget()

    with st.form("s1"):
        c1, c2 = st.columns([2, 1])
        with c1:
            title = st.text_input("프로젝트 제목 (가제)", placeholder="예: 만물트럭 탐정")
        with c2:
            genre = st.selectbox(
                "장르",
                ["미지정", "범죄/스릴러", "드라마", "액션", "로맨스", "코미디",
                 "호러/공포", "SF", "판타지", "코지 미스터리", "느와르", "사회파", "직접 입력"]
            )
            if genre == "직접 입력":
                genre = st.text_input("장르 직접 입력", key="genre_direct")
        
        c3, c4 = st.columns(2)
        with c3:
            target_market = st.selectbox(
                "타겟 시장",
                [
                    "한국 + 글로벌",
                    "한국 (국내)",
                    "일본 (인디·공동제작·리메이크 트랙)",
                    "인도네시아 (JAFF 트랙)",
                    "인도네시아 + 한국 OTT (Netflix SEA)",
                    "한국 + 일본 공동제작",
                    "글로벌 (해외)",
                    "직접 입력",
                ],
                help="Market Lens가 자동 적용됩니다. 일본은 인디·공동제작·리메이크 3트랙만 진입 가능 (외 0점 처리)."
            )
            if target_market == "직접 입력":
                target_market = st.text_input("타겟 시장 직접 입력", key="market_direct")
        with c4:
            format_pref = st.selectbox(
                "선호 포맷",
                ["미정 (Idea Engine이 추천)", "장편 영화", "OTT 시리즈", "미니시리즈",
                 "숏폼 드라마", "웹소설", "웹툰"]
            )
        
        raw_idea = st.text_area(
            "원본 아이디어 (필수)",
            height=220,
            placeholder=(
                "예시:\n"
                "만물트럭(한국) - 이동편의점(일본) 결합\n"
                "만물트럭 탐정 — 셜록홈즈 탐정물\n"
                "고령화된 마을을 돌아다니며 사건 해결\n"
                "추리소설 광이었던 주인공이 깨어나보니 만물트럭 운전사가 되어있었다\n"
                "1. 사라진 시체 / 2. 독극물 살인사건 / 3. 아무도 죽이지 않았다"
            ),
        )
        
        submitted = st.form_submit_button("진단 시작 →", type="primary", use_container_width=True)
        
        if submitted:
            if not title.strip() or not raw_idea.strip():
                st.error("제목과 원본 아이디어는 필수입니다.")
            else:
                st.session_state["stage_1_input"] = {
                    "title": title.strip(),
                    "genre": genre if genre != "미지정" else "장르 미정",
                    "target_market": target_market,
                    "format": format_pref,
                    "raw_idea": raw_idea.strip(),
                }
                st.session_state["current_stage"] = 2
                st.rerun()


def page_stage_2():
    section_header("✍ STEP 2 · 로그라인 정제", "REFINE TO STANDARD")
    small_meta("원본 아이디어를 산업 표준 로그라인 3개 변형으로 정제합니다. Sonnet이 작성합니다.")
    
    inp = st.session_state["stage_1_input"]
    
    if not st.session_state.get("stage_2_logline"):
        if st.button("🪄 로그라인 정제 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            if not client:
                st.warning("ANTHROPIC_API_KEY가 설정되지 않았습니다.")
                return
            with st.spinner("Sonnet이 로그라인 3개 변형 작성 중..."):
                prompt_text = P.LOGLINE_REFINE_PROMPT.format(**inp)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패. 다시 시도해주세요.")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                    return
                st.session_state["stage_2_logline"] = result
                st.rerun()
    else:
        ll = st.session_state["stage_2_logline"]
        for v in ll.get("logline_variants", []):
            st.markdown(f"**[{v['variant']}안 — {v['label']}]**")
            st.markdown(f"<div class='callout'>{v['logline']}</div>", unsafe_allow_html=True)
            cs, cw = st.columns(2)
            with cs:
                st.caption(f"✓ 강점: {v['strength']}")
            with cw:
                st.caption(f"✗ 약점: {v['weakness']}")
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        
        st.success(f"**▶ 추천: {ll.get('recommended', '')}안** — {ll.get('recommendation_reason', '')}")
        st.markdown("---")
        st.markdown("**채택할 로그라인을 선택하세요**")
        
        options = {f"{v['variant']}안 - {v['label']}": v['logline']
                   for v in ll.get("logline_variants", [])}
        recommended_key = None
        for k in options:
            if k.startswith(ll.get("recommended", "A")):
                recommended_key = k
                break
        
        choice = st.radio(
            "로그라인 선택",
            list(options.keys()),
            index=list(options.keys()).index(recommended_key) if recommended_key else 0,
            label_visibility="collapsed"
        )
        st.session_state["selected_logline"] = options[choice]
        
        cb, cr, cn = st.columns([1, 1, 2])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 1
                st.rerun()
        with cr:
            if st.button("재실행"):
                st.session_state["stage_2_logline"] = None
                st.rerun()
        with cn:
            if st.button("Hook 진단으로 →", type="primary", use_container_width=True):
                st.session_state["current_stage"] = 3
                st.rerun()

        # ── 진행 상태 백업 ──
        st.markdown("---")
        render_progress_save_button(stage_num=2)


def page_stage_3():
    """Stage 3 격상판 (v1.2) — Stanton 5원칙 → Hook & Punch 발굴 → 5축 채점"""
    section_header("🎯 STEP 3 · 후크 진단", "FOUNDATION → HOOK & PUNCH → SCORING")
    small_meta(
        "Andrew Stanton 5원칙으로 본질 진단 → Hook & Punch 발굴 → 5축 채점의 3단 구조. "
        "한국 + 할리우드 좌표가 자동 매핑됩니다."
    )

    inp = st.session_state["stage_1_input"]
    logline = st.session_state.get("selected_logline", "")

    # 진행 상태 표시 (3단 stepper)
    foundation_done = bool(st.session_state.get("stage_3_foundation"))
    hp_built_done = bool(st.session_state.get("stage_3_hook_punch_built"))
    scoring_done = bool(st.session_state.get("stage_3_hook"))

    cstep_a, cstep_b, cstep_c = st.columns(3)
    with cstep_a:
        st.markdown(
            f"""<div style="text-align:center;padding:8px;background:{'#FFCB05' if foundation_done else '#F0F2FF'};border-radius:8px;font-weight:700;color:#191970;">
            {'✓' if foundation_done else '①'} 3-A · Stanton 5원칙
            </div>""", unsafe_allow_html=True)
    with cstep_b:
        st.markdown(
            f"""<div style="text-align:center;padding:8px;background:{'#FFCB05' if hp_built_done else '#F0F2FF'};border-radius:8px;font-weight:700;color:#191970;">
            {'✓' if hp_built_done else '②'} 3-B · Hook & Punch
            </div>""", unsafe_allow_html=True)
    with cstep_c:
        st.markdown(
            f"""<div style="text-align:center;padding:8px;background:{'#FFCB05' if scoring_done else '#F0F2FF'};border-radius:8px;font-weight:700;color:#191970;">
            {'✓' if scoring_done else '③'} 3-C · 5축 채점
            </div>""", unsafe_allow_html=True)

    st.markdown("---")

    # ════════════════════════════════════════════════════════
    # 3-A: Stanton 5원칙 진단
    # ════════════════════════════════════════════════════════
    if not foundation_done:
        st.markdown("### 🎬 3-A · Stanton 5원칙 진단")
        st.markdown(
            '<div class="callout">'
            'Andrew Stanton (Pixar)의 스토리텔링 5원칙으로 본질을 진단합니다. '
            'Empathy Anchor · Desire Engine · Stakes Calibration · Emotional Impact · Satisfactory Ending — '
            '각 원칙마다 한국·할리우드 좌표 작품이 자동 매핑됩니다.'
            '</div>', unsafe_allow_html=True)

        if st.button("🎬 Stanton 5원칙 진단 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            if not client:
                st.warning("ANTHROPIC_API_KEY가 설정되지 않았습니다.")
                return
            with st.spinner("Sonnet이 Stanton 5원칙 + 한국·할리우드 좌표 매핑 중... (30~60초)"):
                prompt_text = P.STAGE_3A_STORY_FOUNDATION_PROMPT.format(
                    title=inp["title"], logline=logline,
                    genre=inp["genre"], target_market=inp["target_market"],
                    format=inp["format"], raw_idea=inp["raw_idea"],
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                    return
                st.session_state["stage_3_foundation"] = result
                st.rerun()

        # 이전 / 백업
        st.markdown("---")
        cb, _ = st.columns([1, 3])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 2
                st.rerun()
        render_progress_save_button(stage_num=3)
        return

    # 3-A 결과 표시
    foundation = st.session_state["stage_3_foundation"]
    _render_foundation_result(foundation)

    # ════════════════════════════════════════════════════════
    # 3-B: Hook & Punch 발굴 (질문지 생성 + 답변 + 빌드)
    # ════════════════════════════════════════════════════════
    if not hp_built_done:
        st.markdown("---")
        st.markdown("### 🎯 3-B · Hook & Punch 발굴")

        # 3-B 질문지 생성 (한 번만)
        hp_questions = st.session_state.get("stage_3_hook_punch_questions")
        if not hp_questions:
            st.markdown(
                '<div class="callout">'
                '5원칙 진단을 바탕으로, 이 작품의 <b>Hook(한 줄 후크)</b>과 <b>Punch(잊을 수 없는 한 장면)</b>를 발굴합니다. '
                '10개 질문에 답해주시면 Hook Signature와 Punch Scene이 빌드됩니다.'
                '</div>', unsafe_allow_html=True)

            if st.button("📝 Hook & Punch 질문지 생성 (Sonnet)", type="primary", use_container_width=True):
                client = get_anthropic_client()
                with st.spinner("Sonnet이 Hook 5문 + Punch 5문 생성 중... (20~40초)"):
                    prompt_text = P.STAGE_3B_HOOK_PUNCH_PROMPT.format(
                        title=inp["title"], logline=logline, raw_idea=inp["raw_idea"],
                        foundation_result=json.dumps(foundation, ensure_ascii=False, indent=2),
                    )
                    result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                    if result.get("_parse_error"):
                        st.error("응답 파싱 실패")
                        with st.expander("Raw 응답"):
                            st.text(result.get("_raw", ""))
                        return
                    st.session_state["stage_3_hook_punch_questions"] = result
                    st.rerun()

            _stage3_back_buttons()
            return

        # 질문지 표시 + 답변 입력
        echo = hp_questions.get("echo_back", "")
        if echo:
            st.markdown(f'<div class="callout"><b>진단 종합:</b> {echo}</div>', unsafe_allow_html=True)

        # Reference Hints 표시
        hints = hp_questions.get("reference_hints", {})
        if hints:
            with st.expander("💡 참고 좌표 (한국 + 할리우드)", expanded=False):
                st.markdown(f"**Hook 좌표 — 한국:** {hints.get('hook_reference_korean', '')}")
                st.markdown(f"**Hook 좌표 — 할리우드:** {hints.get('hook_reference_hollywood', '')}")
                st.markdown(f"**Punch 좌표 — 한국:** {hints.get('punch_reference_korean', '')}")
                st.markdown(f"**Punch 좌표 — 할리우드:** {hints.get('punch_reference_hollywood', '')}")

        st.markdown("#### 🎣 Hook 발굴 5문")
        hook_intro = hp_questions.get("hook_extraction_intro", "")
        if hook_intro:
            st.caption(hook_intro)

        hook_answers = {}
        for q in hp_questions.get("hook_questions", []):
            qid = q.get("q_id", "")
            st.markdown(f"**{qid}. {q.get('question', '')}**")
            principle = q.get("principle", "")
            if principle:
                st.caption(f"본질: {principle}")
            hints_q = q.get("hint_options", []) or []
            if hints_q:
                st.caption("보조 옵션: " + " · ".join(hints_q))
            answer = st.text_area(
                f"답변 {qid}", key=f"s3b_{qid}", height=70,
                label_visibility="collapsed",
                placeholder="자유롭게 작성하시거나 보조 옵션 중에서 선택하세요.",
            )
            hook_answers[qid] = answer
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

        st.markdown("#### 💥 Punch 발굴 5문")
        punch_intro = hp_questions.get("punch_extraction_intro", "")
        if punch_intro:
            st.caption(punch_intro)

        punch_answers = {}
        for q in hp_questions.get("punch_questions", []):
            qid = q.get("q_id", "")
            st.markdown(f"**{qid}. {q.get('question', '')}**")
            principle = q.get("principle", "")
            if principle:
                st.caption(f"본질: {principle}")
            hints_q = q.get("hint_options", []) or []
            if hints_q:
                st.caption("보조 옵션: " + " · ".join(hints_q))
            answer = st.text_area(
                f"답변 {qid}", key=f"s3b_{qid}", height=70,
                label_visibility="collapsed",
                placeholder="자유롭게 작성하시거나 보조 옵션 중에서 선택하세요.",
            )
            punch_answers[qid] = answer
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

        st.markdown("---")
        if st.button("🎯 Hook Signature + Punch Scene 빌드 (Sonnet)", type="primary", use_container_width=True):
            # 답변 통합
            all_answers = {**hook_answers, **punch_answers}
            empty = [k for k, v in all_answers.items() if not v.strip()]
            if empty:
                st.warning(f"미답변: {', '.join(empty)} — 10문 모두 답해주세요.")
                return

            st.session_state["stage_3_hook_punch_answers"] = all_answers
            client = get_anthropic_client()
            with st.spinner("Sonnet이 Hook Signature + Punch Scene 빌드 중... (30~50초)"):
                prompt_text = P.STAGE_3B_BUILD_PROMPT.format(
                    title=inp["title"], logline=logline, raw_idea=inp["raw_idea"],
                    foundation_result=json.dumps(foundation, ensure_ascii=False, indent=2),
                    hook_punch_answers=json.dumps(all_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                    return
                st.session_state["stage_3_hook_punch_built"] = result
                st.rerun()

        _stage3_back_buttons()
        return

    # 3-B 결과 표시
    hp_built = st.session_state["stage_3_hook_punch_built"]
    _render_hook_punch_result(hp_built)

    # ════════════════════════════════════════════════════════
    # 3-C: 5축 채점 (기존 Hook Diagnostic)
    # ════════════════════════════════════════════════════════
    if not scoring_done:
        st.markdown("---")
        st.markdown("### 📊 3-C · 5축 채점")
        st.markdown(
            '<div class="callout">'
            '발굴된 Hook & Punch를 바탕으로 5축(구체성 · 갈등 가시성 · 장르 명확성 · 판돈 · 독창성) 채점을 진행합니다. '
            '발굴 단계를 거쳐서 채점 정확도가 향상됩니다.'
            '</div>', unsafe_allow_html=True)

        if st.button("📊 5축 채점 실행 (Sonnet)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 5축 채점 중... (20~30초)"):
                # 채점에 발굴된 Hook을 함께 전달 (정확도 ↑)
                hook_one_liner = hp_built.get("hook_signature", {}).get("hook_one_liner", "")
                enriched_idea = inp["raw_idea"]
                if hook_one_liner:
                    enriched_idea += f"\n\n[발굴된 Hook]: {hook_one_liner}"

                prompt_text = P.HOOK_DIAGNOSTIC_PROMPT.format(
                    title=inp["title"], logline=logline,
                    genre=inp["genre"], format=inp["format"], raw_idea=enriched_idea,
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    return
                st.session_state["stage_3_hook"] = result
                st.rerun()

        _stage3_back_buttons()
        return

    # 3-C 결과 표시 — 기존 채점 UI
    hk = st.session_state["stage_3_hook"]
    st.markdown("---")
    st.markdown("### 📊 3-C · 5축 채점 결과")
    _render_scoring_result(hk)

    # 다음 단계 / 이전 / 재실행
    status = hk.get("gate_status", "")
    st.markdown("---")
    cb, cr, cn = st.columns([1, 1, 2])
    with cb:
        if st.button("← 이전"):
            st.session_state["current_stage"] = 2
            st.rerun()
    with cr:
        if st.button("재실행 (3-C만)"):
            st.session_state["stage_3_hook"] = None
            st.rerun()
    with cn:
        if status == "FAIL":
            st.error("🔴 FAIL — Override 시 진행 가능")
            if st.button("⚠ Override하고 Format으로 →", use_container_width=True):
                st.session_state["current_stage"] = 4
                st.rerun()
        else:
            if st.button("Format 추천으로 →", type="primary", use_container_width=True):
                st.session_state["current_stage"] = 4
                st.rerun()

    # 3-A·3-B 재실행 옵션
    with st.expander("🔄 이전 단계 다시 하기", expanded=False):
        cr_a, cr_b = st.columns(2)
        with cr_a:
            if st.button("3-A Stanton 5원칙 재실행", key="rerun_3a"):
                for k in ["stage_3_foundation", "stage_3_hook_punch_questions",
                          "stage_3_hook_punch_answers", "stage_3_hook_punch_built", "stage_3_hook"]:
                    st.session_state[k] = None
                st.rerun()
        with cr_b:
            if st.button("3-B Hook&Punch 재실행", key="rerun_3b"):
                for k in ["stage_3_hook_punch_questions", "stage_3_hook_punch_answers",
                          "stage_3_hook_punch_built", "stage_3_hook"]:
                    st.session_state[k] = None
                st.rerun()

    # ── 진행 상태 백업 ──
    st.markdown("---")
    render_progress_save_button(stage_num=3)


def _stage3_back_buttons():
    """Stage 3 진행 중 사용하는 공통 뒤로/백업 버튼."""
    st.markdown("---")
    cb, _ = st.columns([1, 3])
    with cb:
        if st.button("← 이전"):
            st.session_state["current_stage"] = 2
            st.rerun()
    render_progress_save_button(stage_num=3)


def _render_foundation_result(foundation):
    """3-A Stanton 5원칙 결과 렌더링."""
    st.markdown("### 🎬 3-A · Stanton 5원칙 진단 결과")

    total = foundation.get("foundation_total_score", 0)
    verdict = foundation.get("foundation_verdict", "")
    weakest = foundation.get("weakest_principle", "")
    strongest = foundation.get("strongest_principle", "")

    # 총점 + 판정
    col_score, col_verdict = st.columns([1, 2])
    with col_score:
        st.markdown(f"""
        <div class="metric-tile">
            <div class="num">{total}/50</div>
            <div class="label">FOUNDATION SCORE</div>
        </div>
        """, unsafe_allow_html=True)
    with col_verdict:
        if "GREEN" in verdict.upper() or "40" in verdict:
            st.success(f"🟢 {verdict}")
        elif "YELLOW" in verdict.upper() or "30" in verdict:
            st.warning(f"🟡 {verdict}")
        else:
            st.error(f"🔴 {verdict}")
        st.caption(f"**가장 강함:** {strongest} · **가장 약함:** {weakest}")

    # 5원칙 카드
    principles_meta = [
        ("principle_1_empathy_anchor", "① Empathy Anchor", "감정이입 가능한 누군가"),
        ("principle_2_desire_engine", "② Desire Engine", "간절히 원하는 것"),
        ("principle_3_stakes_calibration", "③ Stakes Calibration", "어렵지만 가능"),
        ("principle_4_emotional_impact", "④ Emotional Impact", "최대 감정 충격"),
        ("principle_5_satisfactory_ending", "⑤ Satisfactory Ending", "만족스러운 결말"),
    ]

    for key, label, sub in principles_meta:
        p = foundation.get(key, {})
        score = p.get("diagnosis_score", 0)
        bg = "#E8F5E9" if score >= 7 else "#FFF8E1" if score >= 5 else "#FFEBEE"
        border = "#2E7D32" if score >= 7 else "#F57F17" if score >= 5 else "#C62828"

        with st.expander(f"{label} — {sub} · {score}/10", expanded=False):
            st.markdown(f"<div style='background:{bg};border-left:3px solid {border};padding:10px;border-radius:6px;margin-bottom:8px;'>"
                       f"<b>진단:</b> {p.get('diagnosis', '')}</div>", unsafe_allow_html=True)

            # 원칙별 메타데이터
            if key == "principle_1_empathy_anchor":
                st.markdown(f"- **타입:** {p.get('anchor_type', '')}")
                st.markdown(f"- **진입점:** {p.get('entry_point', '')}")
                st.markdown(f"- **시청자 거리:** {p.get('audience_distance', '')}")
            elif key == "principle_2_desire_engine":
                st.markdown(f"- **BJND:** {p.get('bjnd_type', '')}")
                st.markdown(f"- **욕망 대상:** {p.get('desire_target', '')}")
                st.markdown(f"- **강도:** {p.get('desire_intensity', '')}")
            elif key == "principle_3_stakes_calibration":
                st.markdown(f"- **외적 장애물:** {p.get('external_obstacle', '')}")
                st.markdown(f"- **내적 장애물:** {p.get('internal_obstacle', '')}")
                st.markdown(f"- **도달 가능성:** {p.get('achievability', '')}")
            elif key == "principle_4_emotional_impact":
                st.markdown(f"- **Hook 가능성:** {p.get('hook_potential', '')}")
                st.markdown(f"- **Punch 가시성:** {p.get('punch_visibility', '')}")
                st.markdown(f"- **주요 감정:** {p.get('primary_emotion', '')}")
                st.markdown(f"- **시청자 거울:** {p.get('audience_mirror', '')}")
            elif key == "principle_5_satisfactory_ending":
                st.markdown(f"- **결말 유형:** {p.get('ending_type', '')}")
                st.markdown(f"- **카타르시스:** {p.get('catharsis_mechanism', '')}")
                st.markdown(f"- **작가 결말 인지:** {'✓' if p.get('writer_knows_ending') else '✗ (미정)'}")

            # 좌표 매핑
            kr_ref = p.get("korean_reference", "")
            hw_ref = p.get("hollywood_reference", "")
            if kr_ref or hw_ref:
                st.markdown("**좌표 매핑:**")
                if kr_ref:
                    st.markdown(f"  🇰🇷 {kr_ref}")
                if hw_ref:
                    st.markdown(f"  🇺🇸 {hw_ref}")

            improvement = p.get("improvement_note", "")
            if improvement:
                st.caption(f"💡 보강 방향: {improvement}")

    # 다음 단계 안내
    next_step = foundation.get("next_step_guidance", "")
    if next_step:
        st.markdown(f'<div class="callout"><b>다음 단계:</b> {next_step}</div>', unsafe_allow_html=True)


def _render_hook_punch_result(hp_built):
    """3-B Hook & Punch 빌드 결과 렌더링."""
    st.markdown("---")
    st.markdown("### 🎯 3-B · Hook Signature + Punch Scene")

    # Hook Signature 카드
    hs = hp_built.get("hook_signature", {})
    st.markdown(f"""
    <div style="border:2px solid #191970;border-radius:12px;padding:18px;background:#F0F2FF;margin-bottom:12px;">
        <div style="display:inline-block;background:#191970;color:#FFCB05;font-size:.75rem;font-weight:700;padding:3px 10px;border-radius:999px;">🎣 HOOK SIGNATURE</div>
        <div style="font-family:'Playfair Display',serif;font-size:1.15rem;font-weight:700;color:#191970;margin-top:10px;line-height:1.4;">
            "{hs.get('hook_one_liner', '')}"
        </div>
        <div style="margin-top:10px;font-size:.85rem;color:#555;">
            <b>메커니즘:</b> {hs.get('mechanism', '')}<br>
            <b>약속:</b> {hs.get('promise', '')}<br>
            <b>차별점:</b> {hs.get('differentiation', '')}
        </div>
        {f'<div style="margin-top:8px;font-size:.78rem;color:#C62828;font-style:italic;">⚠ 약점: {hs.get("weakness", "")}</div>' if hs.get('weakness') else ''}
    </div>
    """, unsafe_allow_html=True)

    # Punch Scene 카드
    ps = hp_built.get("punch_scene", {})
    sig_pot = ps.get("signature_potential", "MEDIUM")
    sig_color = "#2E7D32" if sig_pot == "HIGH" else "#F57F17" if sig_pot == "MEDIUM" else "#999"

    st.markdown(f"""
    <div style="border:2px solid #FFCB05;border-radius:12px;padding:18px;background:#FFFEF5;margin-bottom:12px;">
        <div style="display:inline-block;background:#FFCB05;color:#191970;font-size:.75rem;font-weight:700;padding:3px 10px;border-radius:999px;">💥 PUNCH SCENE</div>
        <div style="margin-top:10px;font-size:.95rem;color:#1A1A2E;line-height:1.6;">
            {ps.get('scene_description', '')}
        </div>
        <div style="margin-top:10px;font-size:.85rem;color:#555;">
            <b>대사 모드:</b> {ps.get('dialogue_mode', '')} ·
            <b>마지막 컷:</b> {ps.get('final_shot', '')}<br>
            <b>정서:</b> {ps.get('primary_emotion', '')} ·
            <b>배치:</b> {ps.get('placement', '')}
        </div>
        <div style="margin-top:8px;font-size:.78rem;color:{sig_color};font-weight:700;">
            ★ 시각적 서명 가능성: {sig_pot}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 정합 진단
    align = hp_built.get("alignment_diagnosis", {})
    align_check = align.get("alignment_check", "")
    if align_check:
        if "정합" == align_check:
            st.success(f"✓ **Hook ↔ Punch 정합** — {align.get('alignment_reasoning', '')}")
        elif "부분" in align_check:
            st.warning(f"🟡 **부분 정합** — {align.get('alignment_reasoning', '')}")
            if align.get("adjustment_needed"):
                st.caption(f"조정 방향: {align['adjustment_needed']}")
        else:
            st.error(f"🔴 **어긋남** — {align.get('alignment_reasoning', '')}")
            if align.get("adjustment_needed"):
                st.caption(f"조정 방향: {align['adjustment_needed']}")

    # 좌표 매핑
    anchor = hp_built.get("anchor_mapping", {})
    if anchor:
        with st.expander("📍 최종 좌표 매핑", expanded=True):
            st.markdown(f"🇰🇷 **한국 좌표:** {anchor.get('korean_anchor', '')}")
            st.markdown(f"🇺🇸 **할리우드 좌표:** {anchor.get('hollywood_anchor', '')}")
            if anchor.get("market_coordinate"):
                st.markdown(f'<div class="callout"><b>시장 좌표:</b> {anchor["market_coordinate"]}</div>', unsafe_allow_html=True)


def _render_scoring_result(hk):
    """3-C 5축 채점 결과 렌더링 (기존 UI 유지)."""
    scores = hk.get("scores", {})
    axis_kr = {
        "specificity": "구체성", "conflict_visibility": "갈등 가시성",
        "genre_clarity": "장르 명확성", "stakes": "판돈", "originality": "독창성"
    }

    categories = list(axis_kr.values())
    values = [scores.get(k, {}).get("score", 0) for k in axis_kr.keys()]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values, theta=categories, fill='toself', name='Hook Score',
        line=dict(color='#191970', width=2),
        fillcolor='rgba(255, 203, 5, 0.4)'
    ))
    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 10], tickfont=dict(size=10)),
            angularaxis=dict(tickfont=dict(size=12, family='Pretendard'))
        ),
        showlegend=False, height=400, margin=dict(l=80, r=80, t=40, b=40),
        paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)'
    )

    cc, cs = st.columns([1, 1])
    with cc:
        st.plotly_chart(fig, use_container_width=True)
    with cs:
        total = hk.get("total_score", 0)
        status = hk.get("gate_status", "")
        st.markdown(f"""
        <div class="metric-tile">
            <div class="num">{total}/50</div>
            <div class="label">TOTAL HOOK SCORE</div>
        </div>
        """, unsafe_allow_html=True)

        if status == "PASS":
            st.success(f"🟢 **{status}** — 35점 이상으로 GO 진행 가능")
        elif status == "CONDITIONAL":
            st.warning(f"🟡 **{status}** — 약점 보강 필요")
        else:
            st.error(f"🔴 **{status}** — 재고 권장")

    st.markdown("**축별 진단**")
    for k, kr in axis_kr.items():
        sc = scores.get(k, {})
        score = sc.get("score", 0)
        cls = "pass" if score >= 7 else "warn" if score >= 5 else "fail"
        st.markdown(f"""
        <div class="score-card {cls}">
            <span class="axis-name">{kr}</span>
            <span class="axis-score">{score}/10</span>
            <div class="axis-comment">{sc.get('comment', '')}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    cs1, cw1 = st.columns(2)
    with cs1:
        st.markdown("**핵심 강점**")
        for s in hk.get("key_strengths", []):
            st.markdown(f"- {s}")
    with cw1:
        st.markdown("**핵심 약점**")
        for w in hk.get("key_weaknesses", []):
            st.markdown(f"- {w}")

    st.markdown("**보강 제안**")
    for s in hk.get("improvement_suggestions", []):
        st.markdown(f"- {s}")


def page_stage_4():
    section_header("📐 STEP 4 · 포맷 + 장르 + 시장 좌표", "FORMAT · GENRE · MARKET POSITION")
    small_meta(
        "5개 포맷 적합도 · 한국 장르 10분류 매핑 · 시장 좌표 4분류(TENTPOLE/MASTERCLASS/STREAMING/GENRE FEATURE)를 "
        "한 번에 판정합니다. v1.3 신규 — 장르와 시장 좌표가 Creator Engine 룰팩 호출에 직접 활용됩니다."
    )

    inp = st.session_state["stage_1_input"]
    logline = st.session_state.get("selected_logline", "")

    if not st.session_state.get("stage_4_format"):
        if st.button("📐 포맷 + 장르 + 시장 좌표 진단 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 포맷 · 장르 · 시장 좌표 동시 판정 중... (30~60초)"):
                # v1.4: Market Lens 동적 주입
                market_lens_text = MLP.get_lens_text(inp.get("target_market", ""))
                prompt_text = P.FORMAT_RECOMMEND_PROMPT.format(
                    title=inp["title"], logline=logline,
                    genre=inp["genre"],
                    target_market=inp.get("target_market", ""),
                    raw_idea=inp["raw_idea"],
                    market_lens=market_lens_text,
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    return
                st.session_state["stage_4_format"] = result
                st.rerun()
    else:
        fm = st.session_state["stage_4_format"]

        # ════════════════════════════════════════════════════
        # v1.3 신규: 장르 매핑 + 시장 좌표 (상단 우선 표시)
        # ════════════════════════════════════════════════════
        genre_map = fm.get("genre_mapping", {})
        market_pos = fm.get("market_position", {})

        if genre_map or market_pos:
            cgenre, cmarket = st.columns(2)

            # 장르 매핑 카드
            with cgenre:
                st.markdown("#### 🎭 장르 분류 (v1.3)")
                primary_g = genre_map.get("primary_genre", "")
                secondary_g = genre_map.get("secondary_genre")
                genre_text = primary_g
                if secondary_g and secondary_g != "null":
                    genre_text = f"{primary_g} × {secondary_g}"

                st.markdown(f"""
                <div style="border:2px solid #191970;border-radius:12px;padding:14px;background:#F0F2FF;margin-bottom:10px;">
                    <div style="font-size:.7rem;color:#191970;font-weight:700;letter-spacing:.05em;">KOREAN GENRE</div>
                    <div style="font-family:'Playfair Display',serif;font-size:1.4rem;font-weight:700;color:#191970;margin-top:4px;">{genre_text}</div>
                    <div style="font-size:.8rem;color:#555;margin-top:8px;line-height:1.5;">{genre_map.get('primary_reasoning', '')}</div>
                </div>
                """, unsafe_allow_html=True)
                if secondary_g and secondary_g != "null" and genre_map.get("secondary_reasoning"):
                    st.caption(f"복합 장르 결합: {genre_map['secondary_reasoning']}")
                if genre_map.get("korean_genre_anchor"):
                    st.caption(f"🇰🇷 {genre_map['korean_genre_anchor']}")

            # 시장 좌표 카드
            with cmarket:
                st.markdown("#### 🎯 시장 좌표 (v1.3)")
                primary_m = market_pos.get("primary_position", "")
                secondary_m = market_pos.get("secondary_position")

                pos_labels = {
                    "TENTPOLE": ("TENTPOLE", "메이저 영화 흥행", "#191970"),
                    "MASTERCLASS": ("MASTERCLASS", "거장 감독 제안", "#7B1FA2"),
                    "STREAMING": ("STREAMING", "OTT 시리즈", "#E91E63"),
                    "GENRE_FEATURE": ("GENRE FEATURE", "중·소형 장르 영화", "#00897B"),
                }
                label, sub, color = pos_labels.get(primary_m, (primary_m, "", "#191970"))

                st.markdown(f"""
                <div style="border:2px solid {color};border-radius:12px;padding:14px;background:#FFFEF5;margin-bottom:10px;">
                    <div style="font-size:.7rem;color:{color};font-weight:700;letter-spacing:.05em;">MARKET POSITION</div>
                    <div style="font-family:'Playfair Display',serif;font-size:1.4rem;font-weight:700;color:{color};margin-top:4px;">{label}</div>
                    <div style="font-size:.75rem;color:#666;font-style:italic;">{sub}</div>
                    <div style="font-size:.8rem;color:#555;margin-top:8px;line-height:1.5;">{market_pos.get('primary_reasoning', '')}</div>
                </div>
                """, unsafe_allow_html=True)

                if secondary_m and secondary_m != "null":
                    sec_label = pos_labels.get(secondary_m, (secondary_m, "", ""))[0]
                    st.caption(f"2순위 좌표: **{sec_label}** — {market_pos.get('secondary_reasoning', '')}")

                buyers = market_pos.get("target_buyers", [])
                if buyers:
                    st.caption("**판매 대상**: " + " / ".join(buyers))

                if market_pos.get("production_implications"):
                    with st.expander("제작 영향", expanded=False):
                        st.markdown(market_pos["production_implications"])

            st.markdown("---")

        # ════════════════════════════════════════════════════
        # 기존: 5개 포맷 적합도
        # ════════════════════════════════════════════════════
        st.markdown("#### 📊 5개 포맷 적합도")
        fs = fm.get("format_scores", {})
        format_kr = {
            "feature_film": "장편 영화", "ott_series": "OTT 시리즈",
            "mini_series": "미니시리즈", "short_form": "숏폼 드라마", "web_novel": "웹소설"
        }

        for k, kr in format_kr.items():
            f = fs.get(k, {})
            score = f.get("score", 0)
            cls = "pass" if score >= 7 else "warn" if score >= 5 else "fail"
            st.markdown(f"""
            <div class="score-card {cls}">
                <span class="axis-name">{kr}</span>
                <span class="axis-score">{score}/10</span>
                <div class="axis-comment">{f.get('reason', '')}</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")
        primary = fm.get("primary_format_detail", {})
        st.markdown(f"### ▶ 1순위 포맷: **{primary.get('format_name', '')}**")

        ic = st.columns(3)
        with ic[0]:
            if primary.get("episode_count"):
                st.metric("회차", primary["episode_count"])
        with ic[1]:
            if primary.get("runtime_per_episode"):
                st.metric("회당 분량", primary["runtime_per_episode"])
        with ic[2]:
            if primary.get("total_runtime"):
                st.metric("전체 분량", primary["total_runtime"])

        st.markdown("**IP 빌딩 전략**")
        st.markdown(f'<div class="callout">{fm.get("ip_building_strategy", "")}</div>', unsafe_allow_html=True)

        if fm.get("unsuitable_formats"):
            st.markdown("**부적합 포맷**")
            for u in fm["unsuitable_formats"]:
                st.markdown(f"- **{u['format']}**: {u['reason']}")

        st.markdown("---")
        cb, cr, cn = st.columns([1, 1, 2])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 3
                st.rerun()
        with cr:
            if st.button("재실행"):
                st.session_state["stage_4_format"] = None
                st.rerun()
        with cn:
            if st.button("Reference로 →", type="primary", use_container_width=True):
                st.session_state["current_stage"] = 5
                st.rerun()

        # ── 진행 상태 백업 ──
        st.markdown("---")
        render_progress_save_button(stage_num=4)


def page_stage_5():
    section_header("🔍 STEP 5 · 레퍼런스 매핑", "SIMILAR 5 + DIFFERENTIATION")
    small_meta("유사작 5편 발굴 + 차별점 + 치명적 유사작 검증.")
    
    inp = st.session_state["stage_1_input"]
    logline = st.session_state.get("selected_logline", "")
    fm = st.session_state["stage_4_format"]
    primary_format = fm.get("primary_format_detail", {}).get("format_name", inp["format"])
    
    if not st.session_state.get("stage_5_reference"):
        if st.button("🔍 Reference 매핑 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 유사작 5편 분석 중..."):
                # v1.4: Market Lens 동적 주입
                market_lens_text = MLP.get_lens_text(inp.get("target_market", ""))
                prompt_text = P.REFERENCE_MAPPING_PROMPT.format(
                    title=inp["title"], logline=logline,
                    genre=inp["genre"], format=primary_format,
                    target_market=inp.get("target_market", ""),
                    raw_idea=inp["raw_idea"],
                    market_lens=market_lens_text,
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    return
                st.session_state["stage_5_reference"] = result
                st.rerun()
    else:
        rf = st.session_state["stage_5_reference"]
        for ref in rf.get("references", []):
            st.markdown(f"### 《{ref['title']}》")
            st.caption(f"{ref.get('year', '')} · {ref.get('country', '')} · {ref.get('format', '')} · 유사 차원: {ref.get('similarity_axis', '')}")
            st.markdown(f"**공통점**: {ref.get('common_points', '')}")
            st.markdown(f"**차별점**: {ref.get('differentiation', '')}")
            st.markdown("---")
        
        warn = rf.get("lethal_similarity_warning", {})
        if warn.get("exists"):
            st.error(f"⚠ **치명적 유사작 경고**\n\n{warn.get('details', '')}")
        else:
            st.success(f"✓ **치명적 유사작 없음** — {warn.get('details', '안전')}")
        
        st.markdown("**차별화 요약**")
        st.markdown(f'<div class="callout">{rf.get("differentiation_summary", "")}</div>', unsafe_allow_html=True)
        
        st.markdown("**투자자 미팅용 답변**")
        st.markdown(f'> {rf.get("investor_pitch_answer", "")}')
        
        st.markdown("---")
        cb, cr, cn = st.columns([1, 1, 2])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 4
                st.rerun()
        with cr:
            if st.button("재실행"):
                st.session_state["stage_5_reference"] = None
                st.rerun()
        with cn:
            if st.button("Market 진단으로 →", type="primary", use_container_width=True):
                st.session_state["current_stage"] = 6
                st.rerun()

        # ── 진행 상태 백업 ──
        st.markdown("---")
        render_progress_save_button(stage_num=5)


def page_stage_6():
    section_header("📊 STEP 6 · 시장성 진단", "3 MARKET STARS")
    small_meta("한국·글로벌·OTT 3개 시장을 동시에 별점 평가합니다.")
    
    inp = st.session_state["stage_1_input"]
    logline = st.session_state.get("selected_logline", "")
    fm = st.session_state["stage_4_format"]
    primary_format = fm.get("primary_format_detail", {}).get("format_name", inp["format"])
    
    if not st.session_state.get("stage_6_market"):
        if st.button("📊 Market 진단 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 3개 시장 진단 중..."):
                # v1.4: Market Lens 동적 주입
                market_lens_text = MLP.get_lens_text(inp.get("target_market", ""))
                prompt_text = P.MARKET_DIAGNOSTIC_PROMPT.format(
                    title=inp["title"], logline=logline,
                    genre=inp["genre"], primary_format=primary_format,
                    target_market=inp.get("target_market", ""),
                    raw_idea=inp["raw_idea"],
                    market_lens=market_lens_text,
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    return
                st.session_state["stage_6_market"] = result
                st.rerun()
    else:
        mk = st.session_state["stage_6_market"]
        
        # v1.4.1: Market Lens 기반 동적 라벨
        # 우선순위 1: Sonnet 출력의 market_lens_applied
        # 우선순위 2: target_market에서 직접 매핑 (fallback)
        lens_applied = mk.get("market_lens_applied", {})
        primary_market_code = lens_applied.get("primary_market", "")
        if not primary_market_code:
            primary_market_code = MLP.get_primary_market(inp.get("target_market", ""))
        
        # 1차 시장 코드 → 한글 라벨 + 국기 이모지
        market_label_map = {
            "KR": ("한국 시장", "🇰🇷"),
            "JP": ("일본 시장", "🇯🇵"),
            "ID": ("인도네시아 시장", "🇮🇩"),
        }
        primary_label, primary_flag = market_label_map.get(
            primary_market_code, ("한국 시장", "🇰🇷")
        )
        
        c1, c2, c3 = st.columns(3)
        for col, key, label in [
            (c1, "domestic_market", primary_label),
            (c2, "global_market", "글로벌 시장"),
            (c3, "ott_market", "OTT 시장"),
        ]:
            stars = mk.get(key, {}).get("stars", 0)
            with col:
                st.markdown(f"""
                <div class="metric-tile">
                    <div class="num">{'★' * stars}</div>
                    <div class="label">{label}</div>
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        dom = mk.get("domestic_market", {})
        # domestic_market.market_name이 있으면 우선 사용 (Sonnet이 채워준 값)
        domestic_market_name = dom.get("market_name", primary_label.replace(" 시장", ""))
        with st.expander(f"{primary_flag} {domestic_market_name} 시장 (Domestic)", expanded=True):
            ta = dom.get("target_audience", {})
            st.markdown(f"**타겟**: {ta.get('gender', '')} {ta.get('age_range', '')}")
            st.markdown(f"  └ {ta.get('psychographic', '')}")
            st.markdown(f"**예산**: {dom.get('budget_estimate', '')}")
            st.markdown(f"**유통**: {', '.join(dom.get('distribution', []))}")
            st.markdown(f"**IP 확장**: {', '.join(dom.get('ip_extension_potential', []))}")
        
        # v1.4.1: market_lens_applied 정보 표시 (있을 때)
        if lens_applied:
            with st.expander("🎯 Market Lens 적용 정보", expanded=False):
                st.markdown(f"**1차 시장**: {market_label_map.get(lens_applied.get('primary_market', ''), ('-', ''))[0]}")
                sec = lens_applied.get("secondary_market")
                if sec:
                    st.markdown(f"**2차 시장**: {market_label_map.get(sec, ('-', ''))[0]}")
                if lens_applied.get("japan_doc_mode"):
                    st.markdown(f"**JP_DOC 모드**: 활성화")
                    track = lens_applied.get("japan_track", "")
                    st.markdown(f"**일본 진입 트랙**: {track}")
                    reasoning = lens_applied.get("japan_track_reasoning", "")
                    if reasoning:
                        st.markdown(f"  └ {reasoning}")
        
        glb = mk.get("global_market", {})
        with st.expander("🌏 글로벌 시장 (International)", expanded=True):
            st.markdown(f"**1차 타겟**: {glb.get('primary_target_country', '')}")
            st.markdown(f"**어필 포인트**: {glb.get('global_appeal_strength', '')}")
            st.markdown(f"**진입 경로**: {', '.join(glb.get('entry_path', []))}")
            st.markdown(f"**약점**: {glb.get('weakness', '')}")
        
        ott = mk.get("ott_market", {})
        with st.expander("📺 OTT 시장 (Platform)", expanded=True):
            fc = ott.get("first_choice_platform", {})
            sc = ott.get("second_choice_platform", {})
            st.markdown(f"**1순위**: {fc.get('name', '')}")
            st.markdown(f"  └ {fc.get('reason', '')}")
            st.markdown(f"**2순위**: {sc.get('name', '')}")
            st.markdown(f"  └ {sc.get('reason', '')}")
            st.markdown(f"**최적 회차**: {ott.get('optimal_episode_count', '')}")
            st.markdown(f"**경쟁 분석**: {ott.get('competition_analysis', '')}")
        
        timing = mk.get("timing_fit", {})
        st.markdown(f"### 시기적 적합성: {'★' * timing.get('score', 0)}")
        st.markdown(f'<div class="callout">{timing.get("reason", "")}</div>', unsafe_allow_html=True)
        
        st.markdown("### ⚠ 위험 신호")
        for r in mk.get("risk_signals", []):
            st.markdown(f"- {r}")
        
        st.markdown("---")
        cb, cr, cn = st.columns([1, 1, 2])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 5
                st.rerun()
        with cr:
            if st.button("재실행"):
                st.session_state["stage_6_market"] = None
                st.rerun()
        with cn:
            if st.button("최종 판정으로 → (Opus)", type="primary", use_container_width=True):
                st.session_state["current_stage"] = 7
                st.rerun()

        # ── 진행 상태 백업 (Stage 7 직전이라 매우 권장) ──
        st.markdown("---")
        st.warning("⚠ Stage 7(Opus 최종 판정) 진행 전 백업을 강력히 권장합니다. 에러 발생 시 이 JSON으로 1~6단계 복원 가능.")
        render_progress_save_button(stage_num=6)


def page_stage_7():
    section_header("⚖ STEP 7 · 최종 판정", "OPUS · GO / CONDITIONAL / NOGO")
    small_meta("Opus 4.7이 6개 진단 결과를 종합하여 최종 판정과 LOCKED 시드 패키지를 확정합니다.")

    inp = st.session_state["stage_1_input"]

    if not st.session_state.get("stage_7_verdict"):
        st.markdown("""
        <div class="callout">
        <b>Opus 4.7 최종 판정</b><br>
        모든 진단 데이터를 종합하여 GO/CONDITIONAL/NOGO 판정을 내리고, Creator Engine 입력용 LOCKED 시드 패키지를 확정합니다.<br>
        <span style="color:#191970;font-weight:600;">v1.1 — Creator Engine v2.5.2 정합:</span> 핵심 결정·음악 규약·시각 모티프·결말 형식·Creator 의제 5개 신규 LOCKED 영역도 함께 산출됩니다.<br>
        소요 시간: 60~90초
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("⚖ Opus 최종 판정 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 종합 판정 중... (60~90초)"):
                prompt_text = P.FINAL_VERDICT_PROMPT.format(
                    title=inp["title"], raw_idea=inp["raw_idea"],
                    logline_data=json.dumps(st.session_state["stage_2_logline"], ensure_ascii=False, indent=2),
                    hook_data=json.dumps(st.session_state["stage_3_hook"], ensure_ascii=False, indent=2),
                    format_data=json.dumps(st.session_state["stage_4_format"], ensure_ascii=False, indent=2),
                    reference_data=json.dumps(st.session_state["stage_5_reference"], ensure_ascii=False, indent=2),
                    market_data=json.dumps(st.session_state["stage_6_market"], ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                    return
                st.session_state["stage_7_verdict"] = result
                st.rerun()
    else:
        vd = st.session_state["stage_7_verdict"]
        verdict = vd.get("final_verdict", "")
        
        if verdict == "GO":
            st.markdown("""
            <div class="verdict-box go">
                <p class="verdict-label" style="color:#2E7D32;">🟢 GO</p>
            </div>
            """, unsafe_allow_html=True)
        elif verdict == "CONDITIONAL":
            st.markdown("""
            <div class="verdict-box cond">
                <p class="verdict-label" style="color:#F9A825;">🟡 CONDITIONAL</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="verdict-box nogo">
                <p class="verdict-label" style="color:#C62828;">🔴 NOGO</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("### 판정 사유")
        st.markdown(f'<div class="callout">{vd.get("verdict_reasoning", "")}</div>', unsafe_allow_html=True)
        
        if vd.get("conditional_requirements"):
            st.markdown("### 충족 조건")
            for c in vd["conditional_requirements"]:
                st.markdown(f"- {c}")
        
        if vd.get("nogo_alternative"):
            st.markdown("### 대안 제시")
            st.warning(vd["nogo_alternative"])
        
        cd, cp = st.columns(2)
        with cd:
            st.markdown("**✓ 확정된 결정**")
            for k in vd.get("key_decisions_made", []):
                st.markdown(f"- {k}")
        with cp:
            st.markdown("**? Creator Engine에서 결정할 것**")
            for q in vd.get("pending_decisions_for_creator", []):
                st.markdown(f"- {q}")
        
        st.markdown("---")
        st.markdown("### 임원 요약 (Executive Summary)")
        st.success(vd.get("executive_summary", ""))
        
        st.markdown("---")
        section_header("🔑 LOCKED 시드 패키지", "LOCKED SEED PACKAGE")
        
        seed = vd.get("locked_seed_package", {})
        
        st.markdown(f"""
        <div class="locked-card">
            <div class="field-label">Project ID</div>
            <div class="field-value">{seed.get('project_id', '')}</div>
            <div class="field-label">Title (KR / EN)</div>
            <div class="field-value">{seed.get('title_kr', '')} / {seed.get('title_en', '')}</div>
            <div class="field-label">Locked Logline</div>
            <div class="field-value">{seed.get('locked_logline', '')}</div>
        </div>
        """, unsafe_allow_html=True)
        
        cg, cf = st.columns(2)
        with cg:
            gn = seed.get("locked_genre", {})
            st.markdown("**Genre**")
            st.markdown(f"- Primary: {gn.get('primary', '')}")
            st.markdown(f"- Secondary: {gn.get('secondary', '')}")
            if gn.get("tertiary"):
                st.markdown(f"- Tertiary: {gn['tertiary']}")
        with cf:
            ft = seed.get("locked_format", {})
            st.markdown("**Format**")
            st.markdown(f"- Primary: {ft.get('primary', '')}")
            if ft.get("episode_count"):
                st.markdown(f"- Episodes: {ft['episode_count']}")
            if ft.get("runtime"):
                st.markdown(f"- Runtime: {ft['runtime']}")
        
        ct, cth = st.columns(2)
        with ct:
            tg = seed.get("locked_target", {})
            st.markdown("**Target**")
            st.markdown(f"- Domestic: {tg.get('domestic', '')}")
            st.markdown(f"- Global: {tg.get('global', '')}")
        with cth:
            th = seed.get("locked_theme", {})
            st.markdown("**Theme**")
            st.markdown(f"- Surface: {th.get('surface', '')}")
            st.markdown(f"- Deep: {th.get('deep', '')}")
        
        ms = seed.get("locked_market_stars", {})
        st.markdown("**Score Summary**")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("Hook", f"{seed.get('locked_hook_score', 0)}/50")
        with c2:
            st.metric("한국", "★" * ms.get("domestic", 0))
        with c3:
            st.metric("글로벌", "★" * ms.get("global", 0))
        with c4:
            st.metric("OTT", "★" * ms.get("ott", 0))
        
        st.markdown("**Risks to Address (Creator Engine 진행 시 다룰 것)**")
        for r in seed.get("locked_risks_to_address", []):
            st.markdown(f"- {r}")

        # ═══════════════════════════════════════════════════════
        # v1.1 — Creator Engine v2.5.2 정합 5개 신규 LOCKED 영역
        # ═══════════════════════════════════════════════════════
        st.markdown("---")
        section_header(
            "🔒 v1.1 신규 LOCKED 영역",
            "CREATOR ENGINE v2.5.2 ABSORPTION KEYS"
        )
        st.caption(
            "Creator Engine v2.5.2가 작품 본질로 절대 보존하는 5개 영역. "
            "「오랜만에」 검증에서 발견된 핵심 모티프 휘발(61%)을 차단하기 위해 도입."
        )

        # ① locked_core_decisions
        core_decisions = seed.get("locked_core_decisions", []) or []
        with st.expander(
            f"① 확정된 핵심 결정 (locked_core_decisions) · {len(core_decisions)}건",
            expanded=bool(core_decisions),
        ):
            if not core_decisions:
                st.caption("이 작품에는 별도로 LOCK된 핵심 결정이 없습니다 (빈 배열).")
            else:
                for d in core_decisions:
                    if isinstance(d, dict):
                        cat = d.get("category", "")
                        rule = d.get("rule", "") or d.get("decision", "")
                        rationale = d.get("rationale", "")
                        if cat:
                            st.markdown(f"**[{cat}]** {rule}")
                        else:
                            st.markdown(f"- {rule}")
                        if rationale:
                            st.caption(f"근거: {rationale}")
                    elif isinstance(d, str):
                        st.markdown(f"- {d}")

        # ② locked_music_rules
        music_rules = seed.get("locked_music_rules", {}) or {}
        with st.expander(
            f"② 음악 사용 규약 (locked_music_rules) · {'있음' if music_rules else '없음'}",
            expanded=bool(music_rules),
        ):
            if not music_rules:
                st.caption("이 작품에는 음악 사용 규약이 없습니다 (빈 객체). 액션·스릴러·호러에서 자주 발생.")
            elif isinstance(music_rules, dict):
                for k, v in music_rules.items():
                    if isinstance(v, list):
                        st.markdown(f"**{k}**")
                        for item in v:
                            st.markdown(f"- {item}")
                    else:
                        st.markdown(f"**{k}**: {v}")
            elif isinstance(music_rules, list):
                for r in music_rules:
                    st.markdown(f"- {r}")
            else:
                st.markdown(str(music_rules))

        # ③ locked_visual_motifs
        visual_motifs = seed.get("locked_visual_motifs", []) or []
        with st.expander(
            f"③ 시각 모티프 (locked_visual_motifs) · {len(visual_motifs)}건",
            expanded=bool(visual_motifs),
        ):
            if not visual_motifs:
                st.caption("이 작품에는 LOCK된 시각 모티프가 없습니다 (빈 배열).")
            else:
                for m in visual_motifs:
                    if isinstance(m, dict):
                        motif = m.get("motif", "") or m.get("name", "")
                        function = m.get("function", "") or m.get("role", "")
                        if motif and function:
                            st.markdown(f"**{motif}** → {function}")
                        elif motif:
                            st.markdown(f"- {motif}")
                    elif isinstance(m, str):
                        st.markdown(f"- {m}")

        # ④ locked_ending_form
        ending_form = seed.get("locked_ending_form", {}) or {}
        with st.expander(
            f"④ 결말 형식 (locked_ending_form) · {'LOCK됨' if ending_form else '미확정'}",
            expanded=bool(ending_form),
        ):
            if not ending_form:
                st.caption("결말 형식이 LOCK되지 않았습니다 (빈 객체). Creator Engine이 결정.")
            elif isinstance(ending_form, dict):
                if ending_form.get("type"):
                    st.markdown(f"**결말 유형**: {ending_form['type']}")
                if ending_form.get("emotional_resolution"):
                    st.markdown(f"**정서적 해소**: {ending_form['emotional_resolution']}")
                if ending_form.get("final_image"):
                    st.markdown(f"**마지막 이미지**: {ending_form['final_image']}")
                if ending_form.get("forbidden"):
                    st.warning(f"**금지 패턴**: {ending_form['forbidden']}")
            else:
                st.markdown(str(ending_form))

        # ⑤ locked_creator_questions
        creator_questions = seed.get("locked_creator_questions", []) or []
        with st.expander(
            f"⑤ Creator Engine 결정 의제 (locked_creator_questions) · {len(creator_questions)}건",
            expanded=bool(creator_questions),
        ):
            if not creator_questions:
                st.caption("Creator Engine이 답할 의제가 없습니다 (빈 배열).")
            else:
                for q in creator_questions:
                    if isinstance(q, dict):
                        question = q.get("question", "")
                        options = q.get("options", []) or []
                        importance = q.get("importance", "")
                        line = f"**{question}**"
                        if importance:
                            badge_color = {
                                "high": "#C62828",
                                "medium": "#F9A825",
                                "low": "#2E7D32",
                            }.get(importance.lower(), "#666")
                            line += (
                                f" <span style='background:{badge_color};color:white;"
                                f"font-size:.7rem;padding:2px 6px;border-radius:4px;"
                                f"margin-left:6px;'>{importance.upper()}</span>"
                            )
                        st.markdown(line, unsafe_allow_html=True)
                        if options:
                            st.caption(f"후보: {' / '.join(str(o) for o in options)}")
                    elif isinstance(q, str):
                        st.markdown(f"- {q}")

        # ────────────────────────────────────────────────────
        # v1.2 신규 4개 LOCKED 영역 (Stanton 5원칙 + Hook&Punch)
        # ────────────────────────────────────────────────────
        # ⑥ locked_empathy_anchor
        empathy = seed.get("locked_empathy_anchor", {}) or {}
        with st.expander(
            f"⑥ Empathy Anchor (locked_empathy_anchor · v1.2) · {'있음' if empathy else '없음'}",
            expanded=bool(empathy),
        ):
            if not empathy:
                st.caption("이 시드는 v1.1 이전 버전입니다 (빈 객체). v1.2 신규 키 미적용.")
            elif isinstance(empathy, dict):
                if empathy.get("anchor_type"):
                    st.markdown(f"**감정이입 유형**: {empathy['anchor_type']}")
                if empathy.get("entry_point"):
                    st.markdown(f"**진입점**: {empathy['entry_point']}")
                if empathy.get("korean_reference"):
                    st.markdown(f"🇰🇷 **한국 좌표**: {empathy['korean_reference']}")
                if empathy.get("hollywood_reference"):
                    st.markdown(f"🇺🇸 **할리우드 좌표**: {empathy['hollywood_reference']}")

        # ⑦ locked_hook_signature
        hook_sig = seed.get("locked_hook_signature", {}) or {}
        with st.expander(
            f"⑦ Hook Signature (locked_hook_signature · v1.2) · {'있음' if hook_sig else '없음'}",
            expanded=bool(hook_sig),
        ):
            if not hook_sig:
                st.caption("Hook Signature가 LOCK되지 않았습니다 (v1.1 이전 시드).")
            elif isinstance(hook_sig, dict):
                if hook_sig.get("hook_one_liner"):
                    st.markdown(f'<div style="background:#F0F2FF;border-left:3px solid #191970;padding:10px;font-weight:600;font-size:1.05rem;">"{hook_sig["hook_one_liner"]}"</div>', unsafe_allow_html=True)
                if hook_sig.get("mechanism"):
                    st.markdown(f"**메커니즘**: {hook_sig['mechanism']}")
                if hook_sig.get("promise"):
                    st.markdown(f"**약속**: {hook_sig['promise']}")
                if hook_sig.get("differentiation"):
                    st.markdown(f"**차별점**: {hook_sig['differentiation']}")

        # ⑧ locked_punch_scene
        punch = seed.get("locked_punch_scene", {}) or {}
        with st.expander(
            f"⑧ Punch Scene (locked_punch_scene · v1.2) · {'있음' if punch else '없음'}",
            expanded=bool(punch),
        ):
            if not punch:
                st.caption("Punch Scene이 LOCK되지 않았습니다 (v1.1 이전 시드).")
            elif isinstance(punch, dict):
                if punch.get("scene_description"):
                    st.markdown(f'<div style="background:#FFFEF5;border-left:3px solid #FFCB05;padding:10px;font-size:.95rem;line-height:1.6;">{punch["scene_description"]}</div>', unsafe_allow_html=True)
                meta_cols = []
                if punch.get("dialogue_mode"):
                    meta_cols.append(f"**대사 모드**: {punch['dialogue_mode']}")
                if punch.get("final_shot"):
                    meta_cols.append(f"**마지막 컷**: {punch['final_shot']}")
                if punch.get("primary_emotion"):
                    meta_cols.append(f"**정서**: {punch['primary_emotion']}")
                if punch.get("placement"):
                    meta_cols.append(f"**배치**: {punch['placement']}")
                for line in meta_cols:
                    st.markdown(line)
                if punch.get("signature_potential"):
                    sp = punch["signature_potential"]
                    color = {"HIGH": "#2E7D32", "MEDIUM": "#F9A825", "LOW": "#999"}.get(sp, "#666")
                    st.markdown(f"<span style='color:{color};font-weight:700;'>★ 시각적 서명 가능성: {sp}</span>", unsafe_allow_html=True)

        # ⑨ locked_ending_promise
        ending_promise = seed.get("locked_ending_promise", {}) or {}
        with st.expander(
            f"⑨ Ending Promise (locked_ending_promise · v1.2) · {'있음' if ending_promise else '없음'}",
            expanded=bool(ending_promise),
        ):
            if not ending_promise:
                st.caption("결말 약속이 LOCK되지 않았습니다 (v1.1 이전 시드).")
            elif isinstance(ending_promise, dict):
                if ending_promise.get("ending_type"):
                    st.markdown(f"**결말 유형**: {ending_promise['ending_type']}")
                if ending_promise.get("catharsis_mechanism"):
                    st.markdown(f"**카타르시스 메커니즘**: {ending_promise['catharsis_mechanism']}")
                if ending_promise.get("writer_intent"):
                    st.markdown(f"**작가 의도**: {ending_promise['writer_intent']}")
                if ending_promise.get("satisfactory_logic"):
                    st.markdown(f"**만족 결말 논리**: {ending_promise['satisfactory_logic']}")

        # ────────────────────────────────────────────────────
        # v1.3 신규 2개 LOCKED 영역 (장르 + 시장 좌표)
        # ────────────────────────────────────────────────────
        # ⑩ locked_genre_primary
        genre_primary = seed.get("locked_genre_primary", {}) or {}
        with st.expander(
            f"⑩ 장르 분류 (locked_genre_primary · v1.3) · {'있음' if genre_primary else '없음'}",
            expanded=bool(genre_primary),
        ):
            if not genre_primary:
                st.caption("v1.2 이전 시드입니다 (빈 객체). v1.3 신규 키 미적용.")
            elif isinstance(genre_primary, dict):
                primary_g = genre_primary.get("primary", "")
                secondary_g = genre_primary.get("secondary")
                genre_text = primary_g
                if secondary_g and secondary_g not in ("null", None):
                    genre_text = f"{primary_g} × {secondary_g}"
                st.markdown(f'<div style="background:#F0F2FF;border-left:3px solid #191970;padding:10px;font-weight:600;font-size:1.05rem;">한국 장르 매핑: <b>{genre_text}</b></div>', unsafe_allow_html=True)
                if genre_primary.get("reasoning"):
                    st.markdown(f"**근거**: {genre_primary['reasoning']}")
                if genre_primary.get("korean_genre_anchor"):
                    st.caption(f"🇰🇷 {genre_primary['korean_genre_anchor']}")

        # ⑪ locked_market_position
        market_position = seed.get("locked_market_position", {}) or {}
        with st.expander(
            f"⑪ 시장 좌표 (locked_market_position · v1.3) · {'있음' if market_position else '없음'}",
            expanded=bool(market_position),
        ):
            if not market_position:
                st.caption("v1.2 이전 시드입니다 (빈 객체). v1.3 신규 키 미적용.")
            elif isinstance(market_position, dict):
                primary_m = market_position.get("primary", "")
                secondary_m = market_position.get("secondary")
                pos_color = {
                    "TENTPOLE": "#191970",
                    "MASTERCLASS": "#7B1FA2",
                    "STREAMING": "#E91E63",
                    "GENRE_FEATURE": "#00897B",
                }.get(primary_m, "#191970")
                st.markdown(f'<div style="background:#FFFEF5;border-left:3px solid {pos_color};padding:10px;font-weight:600;font-size:1.05rem;color:{pos_color};">시장 좌표: <b>{primary_m}</b></div>', unsafe_allow_html=True)
                if secondary_m and secondary_m not in ("null", None):
                    st.caption(f"2순위 대안: **{secondary_m}**")
                if market_position.get("reasoning"):
                    st.markdown(f"**근거**: {market_position['reasoning']}")
                buyers = market_position.get("target_buyers", [])
                if buyers:
                    st.markdown(f"**판매 대상**: {' / '.join(buyers)}")
                if market_position.get("production_implications"):
                    st.markdown(f"**제작 영향**: {market_position['production_implications']}")

        st.markdown("---")
        section_header("⬇ STEP 8 · 다운로드", "EXPORT")
        
        d1, d2 = st.columns(2)
        with d1:
            docx_bytes = build_diagnostic_docx(dict(st.session_state))
            st.download_button(
                label="📄 진단보고서 DOCX 다운로드",
                data=docx_bytes,
                file_name=f"IdeaDiagnostic_{inp['title']}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with d2:
            json_str = build_seed_json(dict(st.session_state))
            st.download_button(
                label="🔑 LOCKED 시드 JSON 다운로드",
                data=json_str.encode("utf-8"),
                file_name=f"IdeaSeed_{seed.get('project_id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                use_container_width=True,
            )
        
        st.info("📌 **Creator Engine 사용법** — Creator Engine ① 화면 상단의 'Idea Engine JSON 업로드' 버튼을 눌러 위 JSON 파일을 업로드하면 ① 입력 필드가 자동으로 채워집니다.")
        
        with st.expander("JSON 미리보기"):
            st.code(json_str, language="json")
        
        # ── 진행 상태 전체 백업 (Stage 1~7 통째 JSON) ──
        with st.expander("💾 전체 진행 상태 백업 (Stage 1~7 JSON)", expanded=False):
            st.caption("전체 진단 데이터를 JSON으로 보존. 디버깅·재현·아카이브용.")
            render_progress_save_button(stage_num=7)
        
        st.markdown("---")
        cb, cr = st.columns([1, 1])
        with cb:
            if st.button("← 이전"):
                st.session_state["current_stage"] = 6
                st.rerun()
        with cr:
            if st.button("🔄 새 프로젝트 시작", use_container_width=True):
                reset_session()
                st.rerun()


# ═══════════════════════════════════════════════════════════
# v2.0 — HOME PAGE
# ═══════════════════════════════════════════════════════════
def page_home():
    """모드 선택 진입 화면. HUNTER vs TRIAGE 두 카드 제시."""
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown(
        '<div class="callout">'
        '<b>Idea Engine v2.0</b>은 두 개의 트랙으로 구성됩니다. '
        '머릿속에 아직 아이디어가 없다면 <b>HUNTER</b>로 발굴하시고, '
        '이미 아이디어가 있다면 <b>TRIAGE</b>로 진단하세요. '
        'HUNTER에서 발굴한 시드는 자동으로 TRIAGE Stage 1로 전달됩니다.'
        '</div>',
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown("""
        <div style="border:2px solid #E2E2E0;border-radius:14px;padding:28px 26px;background:#fff;height:100%;">
            <div style="display:inline-block;background:#FFCB05;color:#191970;font-size:.72rem;font-weight:700;padding:4px 10px;border-radius:999px;letter-spacing:.05em;margin-bottom:12px;">HUNTER · 발굴</div>
            <div style="font-family:'Playfair Display',serif;font-size:1.8rem;font-weight:700;color:#191970;margin-bottom:6px;">아이디어가 아직 없을 때</div>
            <div style="color:#6B6B7A;font-size:.95rem;margin-bottom:18px;line-height:1.5;">5개의 입구 중 하나로 들어가 작가 안에 잠재된 답을 끌어냅니다.</div>
            <div style="color:#1A1A2E;font-size:.9rem;line-height:1.85;">
                <b>입구 1</b> — 욕망 ("로맨스 만들고 싶다")<br>
                <b>입구 2</b> — 시대 ("IMF 때 이야기")<br>
                <b>입구 3</b> — 트렌드 ("회빙환 해야 하나")<br>
                <b>입구 4</b> — What if ("로또+일주일 루프")<br>
                <b>입구 5</b> — 사실 ("1945.8.15. 일본인")<br>
                <b>입구 0</b> — 자유 텍스트 (자동 분류)
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
        if st.button("🎯 HUNTER 트랙 시작", key="home_to_hunter", type="primary", use_container_width=True):
            st.session_state["mode"] = "HUNTER"
            st.session_state["hunter_entry"] = None
            st.rerun()

    with col2:
        st.markdown("""
        <div style="border:2px solid #E2E2E0;border-radius:14px;padding:28px 26px;background:#fff;height:100%;">
            <div style="display:inline-block;background:#191970;color:#FFCB05;font-size:.72rem;font-weight:700;padding:4px 10px;border-radius:999px;letter-spacing:.05em;margin-bottom:12px;">TRIAGE · 진단</div>
            <div style="font-family:'Playfair Display',serif;font-size:1.8rem;font-weight:700;color:#191970;margin-bottom:6px;">이미 아이디어가 있을 때</div>
            <div style="color:#6B6B7A;font-size:.95rem;margin-bottom:18px;line-height:1.5;">7단계 진단으로 GO/CONDITIONAL/NOGO 판정 + LOCKED 시드 패키지 생성.</div>
            <div style="color:#1A1A2E;font-size:.9rem;line-height:1.85;">
                <b>1</b> 아이디어 입력<br>
                <b>2</b> 로그라인 정제<br>
                <b>3</b> Hook 진단 (Gate 0)<br>
                <b>4</b> Format 추천<br>
                <b>5</b> Reference 매핑<br>
                <b>6</b> Market 진단<br>
                <b>7</b> 최종 판정 (Opus 4.7)
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
        if st.button("🔍 TRIAGE 트랙 시작", key="home_to_triage", use_container_width=True):
            st.session_state["mode"] = "TRIAGE"
            st.rerun()

    st.markdown("<div style='height:24px'></div>", unsafe_allow_html=True)
    st.markdown(
        '<div style="text-align:center;color:#8E8E99;font-size:.82rem;font-family:Pretendard,sans-serif;">'
        '"카탈로그를 보여주는 게 아니라 작가 안에 잠재된 답을 끌어내기"'
        '</div>',
        unsafe_allow_html=True,
    )


# ═══════════════════════════════════════════════════════════
# v2.0 — HUNTER PAGES (골격 — 2~4단계에서 채움)
# ═══════════════════════════════════════════════════════════
def page_hunter_select():
    """입구 선택 화면 — 입구 0 자유 텍스트 + 입구 1~5 카드."""
    section_header("🎯 HUNTER · 입구 선택", "CHOOSE YOUR ENTRY")
    small_meta(
        "5개 입구 중 하나로 들어가시거나, 입구 0에 자유롭게 입력하시면 자동 분류됩니다. "
        "각 입구는 작가의 영감 유형에 맞춘 사고 확장 엔진입니다."
    )

    # ── 입구 0 — 자유 텍스트 ──
    st.markdown("### 🔮 입구 0 — 자유 텍스트 (자동 분류)")
    st.caption("어느 입구로 갈지 모르겠으면, 머릿속에 떠오른 그대로 입력하세요. 자동으로 분류됩니다.")

    with st.form("hunter_entry_0_form", clear_on_submit=False):
        free_text = st.text_area(
            "자유 입력",
            value=st.session_state.get("hunter_input", ""),
            height=100,
            placeholder="예: 로맨스 만들고 싶다 / IMF 때 이야기 / 회빙환 해야 하나 / 로또+일주일 루프 / 1945.8.15 일본인",
        )
        submitted = st.form_submit_button("🔮 자동 분류 실행", type="primary", use_container_width=True)

    if submitted and free_text.strip():
        st.session_state["hunter_input"] = free_text.strip()
        client = get_anthropic_client()
        with st.spinner("Sonnet이 입력을 분석해 적합한 입구로 분류 중... (10~20초)"):
            prompt_text = P.HUNTER_ENTRY_0_PROMPT.format(free_text=free_text.strip())
            result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
            if result.get("_parse_error"):
                st.error("응답 파싱 실패")
                with st.expander("Raw 응답"):
                    st.text(result.get("_raw", ""))
            else:
                st.session_state["hunter_classified"] = result
                st.rerun()

    # ── 자동 분류 결과 표시 ──
    classified = st.session_state.get("hunter_classified")
    if classified:
        st.markdown("---")
        st.markdown("### 🎯 자동 분류 결과")
        primary = classified.get("primary_entry", {})
        secondary = classified.get("secondary_entry", {})

        st.markdown(f"""
        <div class="callout">
        <b>1순위:</b> 입구 {primary.get('entry_id', '?')} — {primary.get('entry_name', '')}
        (확신도 {primary.get('confidence', 0)}%)<br>
        <i>{primary.get('reasoning', '')}</i>
        </div>
        """, unsafe_allow_html=True)

        if secondary.get("entry_id"):
            st.caption(f"2순위: 입구 {secondary['entry_id']} — {secondary.get('entry_name', '')}: {secondary.get('reasoning', '')}")

        st.markdown(f"**입력 재진술**: {classified.get('restated_input', '')}")

        first_qs = classified.get("first_questions", [])
        if first_qs:
            st.markdown("**1차 사고 확장 질문 (다음 입구에서 답하실 것):**")
            for q in first_qs:
                st.markdown(f"- {q}")

        col_go, col_alt = st.columns(2)
        with col_go:
            if st.button(f"→ 입구 {primary.get('entry_id', '1')}로 진입", key="goto_primary", type="primary", use_container_width=True):
                st.session_state["hunter_entry"] = str(primary.get("entry_id", "1"))
                st.rerun()
        with col_alt:
            if secondary.get("entry_id"):
                if st.button(f"→ 입구 {secondary['entry_id']}로 진입 (2순위)", key="goto_secondary", use_container_width=True):
                    st.session_state["hunter_entry"] = str(secondary["entry_id"])
                    st.rerun()

    # ── 입구 1~5 직접 카드 ──
    st.markdown("---")
    st.markdown("### 🚪 입구를 직접 선택")
    st.caption("어느 입구에 들어갈지 명확하면 아래에서 직접 선택하세요.")

    entries_meta = [
        ("1", "결핍·상실", "LACK & LOSS", "💔",
         "'~을 만들고 싶다'는 갈망에서 출발. BJND 진단으로 결핍/상실 본질 판별."),
        ("2", "시대", "PERIOD", "🕰️",
         "특정 시대(IMF·1990년대 등)에 끌릴 때. 시대 진단 + 디테일 펼침."),
        ("3", "트렌드", "TREND", "📈",
         "현재 시장 트렌드(회빙환·숏폼 등) 추종/변주/회피 결정."),
        ("4", "What if", "HYPOTHESIS", "❓",
         "'만약 ~라면?' 가설 확장 + 4대 함정 경고 + 톤 3분기."),
        ("5", "사실", "FACT", "📜",
         "구체적 역사·실화·뉴스 작품화. 5시점 발굴."),
    ]

    cols = st.columns(5)
    for i, (eid, kr, en, emoji, desc) in enumerate(entries_meta):
        with cols[i]:
            st.markdown(f"""
            <div style="border:1.5px solid #E2E2E0;border-radius:12px;padding:14px 12px;background:#fff;min-height:200px;">
                <div style="font-size:1.8rem;text-align:center;margin-bottom:6px;">{emoji}</div>
                <div style="font-family:'Playfair Display',serif;font-size:.9rem;font-weight:700;color:#191970;text-align:center;margin-bottom:4px;">입구 {eid}</div>
                <div style="font-size:1rem;font-weight:700;color:#1A1A2E;text-align:center;margin-bottom:6px;">{kr}</div>
                <div style="font-size:.65rem;color:#999;text-align:center;letter-spacing:.05em;margin-bottom:10px;">{en}</div>
                <div style="font-size:.72rem;color:#555;line-height:1.4;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
            if st.button(f"입구 {eid} 진입", key=f"direct_entry_{eid}", use_container_width=True):
                st.session_state["hunter_entry"] = eid
                st.rerun()

    # ── 개발자 디버그 ──
    with st.expander("개발자: 현재 HUNTER 상태", expanded=False):
        st.json({
            "mode": st.session_state.get("mode"),
            "hunter_entry": st.session_state.get("hunter_entry"),
            "hunter_input": st.session_state.get("hunter_input"),
            "hunter_classified": st.session_state.get("hunter_classified"),
            "hunter_stage_data_keys": list((st.session_state.get("hunter_stage_data") or {}).keys()),
            "hunter_output": st.session_state.get("hunter_output"),
        })


def _hunter_render_questions(questions, key_prefix, intro_text=None):
    """공통 질문 렌더링 헬퍼. 5개 질문 + 보조 옵션 + 답변 입력."""
    if intro_text:
        st.markdown(f'<div class="callout">{intro_text}</div>', unsafe_allow_html=True)

    answers = {}
    for q in questions:
        qid = q.get("q_id", 0)
        st.markdown(f"**Q{qid}. {q.get('question', '')}**")
        principle = q.get("principle", "")
        if principle:
            st.caption(f"원칙: {principle}")

        hints = q.get("hint_options", []) or []
        if hints:
            st.caption("보조 옵션: " + " · ".join(hints))

        answer = st.text_area(
            f"답변 {qid}",
            key=f"{key_prefix}_q{qid}",
            height=80,
            label_visibility="collapsed",
            placeholder="자유롭게 작성하시거나, 보조 옵션 중 하나를 선택해서 살을 붙이세요.",
        )
        answers[f"q{qid}"] = answer
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    return answers


def _hunter_render_seed_cards(seeds, bjnd_essence=None, period_essence=None):
    """시드 후보 카드 렌더링 + 선택 버튼."""
    selected_seed = None

    for s in seeds:
        sid = s.get("seed_id", "")
        slabel = s.get("seed_label", "")
        title = s.get("title", "")
        genre = s.get("genre", "")
        target = s.get("target_market", "")
        fmt = s.get("format", "")
        raw_idea = s.get("raw_idea", "")
        diff = s.get("differentiation", "")
        bjnd_label = s.get("bjnd_label", "") or s.get("period_bjnd_label", "")

        with st.container():
            st.markdown(f"""
            <div style="border:2px solid #FFCB05;border-radius:14px;padding:18px;background:#FFFEF5;margin-bottom:14px;">
                <div style="display:inline-block;background:#191970;color:#FFCB05;font-size:.7rem;font-weight:700;padding:3px 10px;border-radius:999px;">시드 {sid} · {slabel}</div>
                <div style="font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:700;color:#191970;margin-top:8px;">{title}</div>
                <div style="font-size:.85rem;color:#666;margin-top:4px;">
                    <b>장르</b>: {genre} · <b>포맷</b>: {fmt}<br>
                    <b>타겟</b>: {target}
                </div>
                {f'<div style="font-size:.75rem;color:#191970;font-weight:600;margin-top:6px;">{bjnd_label}</div>' if bjnd_label else ''}
                <div style="margin-top:10px;padding:10px;background:white;border-radius:8px;font-size:.88rem;line-height:1.6;color:#1A1A2E;">{raw_idea}</div>
                <div style="font-size:.78rem;color:#555;margin-top:8px;font-style:italic;">{diff}</div>
            </div>
            """, unsafe_allow_html=True)

            if st.button(f"→ 시드 {sid} 선택해서 TRIAGE로 전송", key=f"select_seed_{sid}", use_container_width=True, type="primary"):
                # HUNTER 시드 → TRIAGE 인계 형식으로 변환
                st.session_state["hunter_output"] = {
                    "title": title,
                    "genre": genre,
                    "target_market": target,
                    "format_pref": fmt,
                    "raw_idea": raw_idea,
                    "hunter_meta": {
                        "entry": st.session_state.get("hunter_entry", ""),
                        "seed_id": sid,
                        "seed_label": slabel,
                        "bjnd_label": bjnd_label,
                        "bjnd_essence": bjnd_essence,
                        "period_essence": period_essence,
                        "differentiation": diff,
                    },
                }
                transfer_hunter_seed_to_triage()
                st.session_state["mode"] = "TRIAGE"
                st.rerun()


def _hunter_back_button(entry_id):
    """입구 선택으로 돌아가기 버튼."""
    col_back, _ = st.columns([1, 4])
    with col_back:
        if st.button("← 입구 선택으로", key=f"hunter_back_{entry_id}", use_container_width=True):
            st.session_state["hunter_entry"] = None
            # 진행 중 데이터는 보존 (작가가 다시 들어올 수 있도록)
            st.rerun()


def _hunter_reset_button(entry_id):
    """현재 입구 진행 데이터 초기화 버튼."""
    if st.button("🔄 이 입구 처음부터 다시", key=f"hunter_reset_{entry_id}", use_container_width=True):
        # 현재 입구 진행 데이터만 삭제
        stage_data = st.session_state.get("hunter_stage_data", {})
        for k in list(stage_data.keys()):
            if k.startswith(f"entry{entry_id}_"):
                del stage_data[k]
        st.session_state["hunter_stage_data"] = stage_data
        st.rerun()


def page_hunter_entry(entry_id: str):
    """입구별 페이지 — 입구 1~5 본 구현."""
    entry_titles = {
        "1": ("결핍·상실", "LACK & LOSS"),
        "2": ("시대", "PERIOD"),
        "3": ("트렌드", "TREND"),
        "4": ("What if", "HYPOTHESIS"),
        "5": ("사실", "FACT"),
    }
    kr, en = entry_titles.get(entry_id, ("입구", "ENTRY"))
    section_header(f"🎯 HUNTER · 입구 {entry_id} — {kr}", en)

    # 입구별 라우팅
    if entry_id == "1":
        _hunter_entry_1_lack_loss()
    elif entry_id == "2":
        _hunter_entry_2_period()
    elif entry_id == "3":
        _hunter_entry_3_trend()
    elif entry_id == "4":
        _hunter_entry_4_whatif()
    elif entry_id == "5":
        _hunter_entry_5_fact()
    else:
        st.warning(f"알 수 없는 입구: {entry_id}")
        _hunter_back_button(entry_id)


# ────────────────────────────────────────────────────────────
# 입구 1 — 결핍·상실 (3턴: 진단 → 정밀 확장 → 시드)
# ────────────────────────────────────────────────────────────
def _hunter_entry_1_lack_loss():
    small_meta(
        "BJND 진단(결핍 vs 상실) → 정밀 사고 확장 → 시드 후보 3개의 3턴 구조입니다. "
        "같은 표면 갈망도 결핍/상실 발생 근원에 따라 완전히 다른 작품이 됩니다."
    )

    stage_data = st.session_state.setdefault("hunter_stage_data", {})

    # ── 턴 1: 입력 받기 ──
    desire_input = stage_data.get("entry1_desire_input", "")
    if not desire_input:
        # 입구 0에서 자동 분류된 입력이 있으면 가져오기
        if st.session_state.get("hunter_classified", {}).get("primary_entry", {}).get("entry_id") == 1:
            desire_input = st.session_state.get("hunter_input", "")

        st.markdown("### 1단계 — 작품을 향한 갈망을 입력하세요")
        with st.form("entry1_desire_form"):
            desire = st.text_area(
                "어떤 작품을 만들고 싶으신가요?",
                value=desire_input,
                height=100,
                placeholder="예: 로맨스 만들고 싶다 / 복수극이 끌려 / 감동적인 가족 드라마 / 미스터리 스릴러",
            )
            submitted = st.form_submit_button("→ BJND 진단 시작", type="primary", use_container_width=True)

        if submitted and desire.strip():
            stage_data["entry1_desire_input"] = desire.strip()
            st.session_state["hunter_stage_data"] = stage_data
            st.rerun()

        _hunter_back_button("1")
        return

    # 입력 표시
    st.markdown(f"**작가의 갈망:** _{desire_input}_")
    st.markdown("---")

    # ── 턴 2: BJND 진단 질문 ──
    diagnosis = stage_data.get("entry1_diagnosis")
    if not diagnosis:
        st.markdown("### 2단계 — BJND 진단 질문 5개")
        if st.button("📊 BJND 진단 질문 생성 (Sonnet 4.6)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 BJND 진단 질문 5개 생성 중... (20~40초)"):
                prompt_text = P.HUNTER_ENTRY_1_DIAGNOSIS_PROMPT.format(desire_input=desire_input)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry1_diagnosis"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()

        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("1")
        with col_reset:
            _hunter_reset_button("1")
        return

    # 진단 질문 표시 + 답변 입력
    diagnosis_answers = stage_data.get("entry1_diagnosis_answers")
    if not diagnosis_answers:
        st.markdown("### 2단계 — BJND 진단 5개 질문에 답해주세요")
        echo = diagnosis.get("echo_back", "")
        if echo:
            st.markdown(f'<div class="callout"><b>입력 재진술:</b> {echo}</div>', unsafe_allow_html=True)

        intro = diagnosis.get("diagnosis_intro", "")
        answers = _hunter_render_questions(
            diagnosis.get("diagnosis_questions", []),
            "e1_diag",
            intro_text=intro,
        )

        if st.button("→ 진단 결과 + 정밀 사고 확장", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry1_diagnosis_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")

        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("1")
        with col_reset:
            _hunter_reset_button("1")
        return

    # ── 턴 3: BJND 판정 + 정밀 확장 질문 ──
    expansion = stage_data.get("entry1_expansion")
    if not expansion:
        st.markdown("### 3단계 — BJND 판정 + 정밀 사고 확장 (Sonnet 4.6)")
        if st.button("🔬 BJND 판정 + 정밀 확장 실행", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 결핍/상실 판정 + 정밀 질문 생성 중... (30~50초)"):
                prompt_text = P.HUNTER_ENTRY_1_EXPANSION_PROMPT.format(
                    desire_input=desire_input,
                    diagnosis_answers=json.dumps(diagnosis_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry1_expansion"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()

        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("1")
        with col_reset:
            _hunter_reset_button("1")
        return

    # 판정 결과 표시
    diag_result = expansion.get("diagnosis_result", {})
    bjnd_type = diag_result.get("type", "")
    lack_score = diag_result.get("lack_score", 0)
    loss_score = diag_result.get("loss_score", 0)

    st.markdown("### 📊 BJND 판정 결과")
    col_t, col_l, col_lo = st.columns(3)
    with col_t:
        st.metric("판정 유형", bjnd_type)
    with col_l:
        st.metric("결핍 점수", f"{lack_score}/5")
    with col_lo:
        st.metric("상실 점수", f"{loss_score}/5")

    st.markdown(f'<div class="callout">{diag_result.get("reasoning", "")}</div>', unsafe_allow_html=True)

    refs = diag_result.get("reference_works", [])
    if refs:
        st.markdown("**참고작 (같은 결의 작품):**")
        for r in refs:
            st.markdown(f"- {r}")

    st.markdown("---")

    # 정밀 확장 질문 답변
    expansion_answers = stage_data.get("entry1_expansion_answers")
    if not expansion_answers:
        st.markdown("### 4단계 — 정밀 사고 확장 5개 질문에 답해주세요")
        intro = expansion.get("expansion_intro", "")
        answers = _hunter_render_questions(
            expansion.get("expansion_questions", []),
            "e1_exp",
            intro_text=intro,
        )

        if st.button("→ 시드 후보 3개 빌드", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry1_expansion_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")

        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("1")
        with col_reset:
            _hunter_reset_button("1")
        return

    # ── 턴 4: 시드 빌드 ──
    seeds_result = stage_data.get("entry1_seeds")
    if not seeds_result:
        st.markdown("### 5단계 — 시드 후보 3개 빌드 (Opus 4.7)")
        if st.button("🌱 시드 후보 3개 생성", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 시드 후보 3개 빌드 중... (40~60초)"):
                prompt_text = P.HUNTER_ENTRY_1_SEEDS_PROMPT.format(
                    desire_input=desire_input,
                    diagnosis_result=json.dumps(diag_result, ensure_ascii=False, indent=2),
                    expansion_answers=json.dumps(expansion_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry1_seeds"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()

        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("1")
        with col_reset:
            _hunter_reset_button("1")
        return

    # 시드 후보 표시 + 선택
    st.markdown("### 🌱 시드 후보 3개")
    synthesis = seeds_result.get("synthesis", "")
    if synthesis:
        st.markdown(f'<div class="callout"><b>본질 종합:</b> {synthesis}</div>', unsafe_allow_html=True)

    bjnd_essence = seeds_result.get("bjnd_essence", {})
    if bjnd_essence:
        with st.expander("BJND 본질 상세"):
            st.json(bjnd_essence)

    _hunter_render_seed_cards(seeds_result.get("seeds", []), bjnd_essence=bjnd_essence)

    recommendation = seeds_result.get("recommendation", "")
    if recommendation:
        st.markdown(f"**추천:** {recommendation}")

    st.caption(seeds_result.get("next_step", ""))

    col_back, col_reset = st.columns(2)
    with col_back:
        _hunter_back_button("1")
    with col_reset:
        _hunter_reset_button("1")


# ────────────────────────────────────────────────────────────
# 입구 2 — 시대 (3턴: 시대 진단 → 디테일 펼침 → 시드)
# ────────────────────────────────────────────────────────────
def _hunter_entry_2_period():
    small_meta(
        "시대 진단(결핍형/상실형) → 시대 디테일 펼침 → 시드 후보 3개의 3턴 구조입니다. "
        "그 시대에 대한 작가의 관계가 결핍(못 가본 시대 동경)인지 상실(잃어버린 시대 회한)인지 진단합니다."
    )

    stage_data = st.session_state.setdefault("hunter_stage_data", {})

    # 턴 1: 입력
    period_input = stage_data.get("entry2_period_input", "")
    if not period_input:
        if st.session_state.get("hunter_classified", {}).get("primary_entry", {}).get("entry_id") == 2:
            period_input = st.session_state.get("hunter_input", "")

        st.markdown("### 1단계 — 어떤 시대를 작품 배경으로 하고 싶으신가요?")
        with st.form("entry2_period_form"):
            period = st.text_area(
                "시대 입력",
                value=period_input,
                height=80,
                placeholder="예: IMF 때 이야기 / 1990년대 한국 / 조선 후기 / 2002년 월드컵 / 식민지 시대",
            )
            submitted = st.form_submit_button("→ 시대 진단 시작", type="primary", use_container_width=True)
        if submitted and period.strip():
            stage_data["entry2_period_input"] = period.strip()
            st.session_state["hunter_stage_data"] = stage_data
            st.rerun()
        _hunter_back_button("2")
        return

    st.markdown(f"**작가의 시대:** _{period_input}_")
    st.markdown("---")

    # 턴 2: 시대 진단 질문
    diagnosis = stage_data.get("entry2_diagnosis")
    if not diagnosis:
        if st.button("📊 시대 진단 질문 생성", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 시대 진단 질문 생성 중... (20~40초)"):
                prompt_text = P.HUNTER_ENTRY_2_DIAGNOSIS_PROMPT.format(period_input=period_input)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry2_diagnosis"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("2")
        with col_reset:
            _hunter_reset_button("2")
        return

    # 진단 질문 답변
    diagnosis_answers = stage_data.get("entry2_diagnosis_answers")
    if not diagnosis_answers:
        echo = diagnosis.get("echo_back", "")
        if echo:
            st.markdown(f'<div class="callout"><b>입력 재진술:</b> {echo}</div>', unsafe_allow_html=True)
        intro = diagnosis.get("diagnosis_intro", "")
        answers = _hunter_render_questions(
            diagnosis.get("diagnosis_questions", []),
            "e2_diag",
            intro_text=intro,
        )
        if st.button("→ 시대 판정 + 디테일 펼침", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry2_diagnosis_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("2")
        with col_reset:
            _hunter_reset_button("2")
        return

    # 턴 3: 시대 판정 + 디테일 펼침
    expansion = stage_data.get("entry2_expansion")
    if not expansion:
        if st.button("🔬 시대 판정 + 디테일 펼침", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 시대 판정 + 디테일 펼침 중... (30~50초)"):
                prompt_text = P.HUNTER_ENTRY_2_EXPANSION_PROMPT.format(
                    period_input=period_input,
                    diagnosis_answers=json.dumps(diagnosis_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry2_expansion"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("2")
        with col_reset:
            _hunter_reset_button("2")
        return

    # 시대 판정 표시
    diag_result = expansion.get("diagnosis_result", {})
    st.markdown("### 🕰️ 시대 판정")
    st.markdown(f"**유형:** {diag_result.get('type', '')}")
    st.markdown(f"**시대 명명:** {diag_result.get('period_label', '')}")
    st.markdown(f'<div class="callout">{diag_result.get("reasoning", "")}</div>', unsafe_allow_html=True)

    canvas = expansion.get("period_detail_canvas", {})
    if canvas:
        with st.expander("📜 시대 디테일 캔버스", expanded=True):
            st.markdown(f"**감각적 정수:** {canvas.get('sensory_essence', '')}")
            st.markdown(f"**사회 풍경:** {canvas.get('social_landscape', '')}")
            st.markdown(f"**결핍/상실 풍경:** {canvas.get('lack_or_loss_landscape', '')}")
            add = canvas.get("additional_details", [])
            if add:
                st.markdown("**추가 디테일:**")
                for d in add:
                    st.markdown(f"- {d}")

    st.markdown("---")

    # 정밀 확장 답변
    expansion_answers = stage_data.get("entry2_expansion_answers")
    if not expansion_answers:
        intro = expansion.get("expansion_intro", "")
        answers = _hunter_render_questions(
            expansion.get("expansion_questions", []),
            "e2_exp",
            intro_text=intro,
        )
        if st.button("→ 시드 후보 3개 빌드", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry2_expansion_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("2")
        with col_reset:
            _hunter_reset_button("2")
        return

    # 턴 4: 시드 빌드
    seeds_result = stage_data.get("entry2_seeds")
    if not seeds_result:
        if st.button("🌱 시드 후보 3개 생성 (Opus 4.7)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 시대 시드 3개 빌드 중... (40~60초)"):
                prompt_text = P.HUNTER_ENTRY_2_SEEDS_PROMPT.format(
                    period_input=period_input,
                    diagnosis_result=json.dumps(diag_result, ensure_ascii=False, indent=2),
                    expansion_answers=json.dumps(expansion_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry2_seeds"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("2")
        with col_reset:
            _hunter_reset_button("2")
        return

    # 시드 표시
    st.markdown("### 🌱 시대 시드 후보 3개")
    synthesis = seeds_result.get("synthesis", "")
    if synthesis:
        st.markdown(f'<div class="callout"><b>본질 종합:</b> {synthesis}</div>', unsafe_allow_html=True)

    period_essence = seeds_result.get("period_essence", {})
    if period_essence:
        with st.expander("시대 본질 상세"):
            st.json(period_essence)

    _hunter_render_seed_cards(seeds_result.get("seeds", []), period_essence=period_essence)

    if seeds_result.get("recommendation"):
        st.markdown(f"**추천:** {seeds_result['recommendation']}")
    st.caption(seeds_result.get("next_step", ""))

    col_back, col_reset = st.columns(2)
    with col_back:
        _hunter_back_button("2")
    with col_reset:
        _hunter_reset_button("2")


# ────────────────────────────────────────────────────────────
# 입구 3 — 트렌드 (2턴: 분석+3길 → 시드)
# ────────────────────────────────────────────────────────────
def _hunter_entry_3_trend():
    small_meta(
        "트렌드 분석 + 추종/변주/회피 3길 → 시드 후보 3개의 2턴 구조입니다. "
        "BJND는 적용하지 않습니다 (장르·포맷 결정 입구이지 인물 욕망 입구가 아니므로)."
    )

    stage_data = st.session_state.setdefault("hunter_stage_data", {})

    # 턴 1: 입력
    trend_input = stage_data.get("entry3_trend_input", "")
    if not trend_input:
        if st.session_state.get("hunter_classified", {}).get("primary_entry", {}).get("entry_id") == 3:
            trend_input = st.session_state.get("hunter_input", "")

        st.markdown("### 1단계 — 어떤 트렌드에 대해 입장을 정리하고 싶으신가요?")
        with st.form("entry3_trend_form"):
            trend = st.text_area(
                "트렌드 입력",
                value=trend_input,
                height=80,
                placeholder="예: 회빙환 해야 하나 / 숏폼 드라마가 대세 / SF가 뜨고 있다 / 사극 부활 / 로맨스 판타지",
            )
            submitted = st.form_submit_button("→ 트렌드 분석 시작", type="primary", use_container_width=True)
        if submitted and trend.strip():
            stage_data["entry3_trend_input"] = trend.strip()
            st.session_state["hunter_stage_data"] = stage_data
            st.rerun()
        _hunter_back_button("3")
        return

    st.markdown(f"**작가의 트렌드:** _{trend_input}_")
    st.markdown("---")

    # 턴 2: 트렌드 분석 + 3길 + 진단 질문
    diagnosis = stage_data.get("entry3_diagnosis")
    if not diagnosis:
        if st.button("📈 트렌드 분석 + 3길 펼침", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 트렌드 분석 + 3길 + 진단 질문 생성 중... (30~50초)"):
                prompt_text = P.HUNTER_ENTRY_3_DIAGNOSIS_PROMPT.format(trend_input=trend_input)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry3_diagnosis"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("3")
        with col_reset:
            _hunter_reset_button("3")
        return

    # 트렌드 분석 표시
    analysis = diagnosis.get("trend_analysis", {})
    st.markdown("### 📈 트렌드 분석")
    st.markdown(f"**본질:** {analysis.get('trend_essence', '')}")
    st.markdown(f"**시장 위치:** {analysis.get('market_position', '')}")
    st.caption(analysis.get("market_reasoning", ""))

    st.markdown("---")
    st.markdown("### 🛤️ 추종 / 변주 / 회피 3길")
    paths = diagnosis.get("three_paths", {})
    col_f, col_v, col_a = st.columns(3)
    for col, key, label in [(col_f, "follow", "추종"), (col_v, "variation", "변주"), (col_a, "avoidance", "회피")]:
        p = paths.get(key, {})
        with col:
            st.markdown(f"**{label}**")
            st.caption(p.get("definition", ""))
            st.markdown(f"<span style='color:#2E7D32;font-size:.8rem;'>장점: {p.get('advantage', '')}</span>", unsafe_allow_html=True)
            st.markdown(f"<span style='color:#C62828;font-size:.8rem;'>위험: {p.get('risk', '')}</span>", unsafe_allow_html=True)
            examples = p.get("examples", [])
            if examples:
                st.caption("예: " + " / ".join(examples))

    st.markdown("---")

    # 진단 답변
    diagnosis_answers = stage_data.get("entry3_diagnosis_answers")
    if not diagnosis_answers:
        intro = "트렌드 분석을 보셨으니 이제 작가 본인의 결을 진단합니다."
        answers = _hunter_render_questions(
            diagnosis.get("diagnosis_questions", []),
            "e3_diag",
            intro_text=intro,
        )
        if st.button("→ 시드 후보 3개 빌드", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry3_diagnosis_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("3")
        with col_reset:
            _hunter_reset_button("3")
        return

    # 시드 빌드
    seeds_result = stage_data.get("entry3_seeds")
    if not seeds_result:
        if st.button("🌱 시드 후보 3개 생성 (Opus 4.7)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 트렌드 노선 시드 3개 빌드 중... (40~60초)"):
                prompt_text = P.HUNTER_ENTRY_3_SEEDS_PROMPT.format(
                    trend_input=trend_input,
                    trend_analysis=json.dumps(analysis, ensure_ascii=False, indent=2),
                    diagnosis_answers=json.dumps(diagnosis_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry3_seeds"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("3")
        with col_reset:
            _hunter_reset_button("3")
        return

    # 시드 표시
    selected_path = seeds_result.get("selected_path", "")
    st.markdown(f"### 🌱 시드 후보 3개 — 선택 노선: **{selected_path}**")
    if seeds_result.get("path_reasoning"):
        st.markdown(f'<div class="callout">{seeds_result["path_reasoning"]}</div>', unsafe_allow_html=True)
    if seeds_result.get("synthesis"):
        st.markdown(f"**본질 종합:** {seeds_result['synthesis']}")

    _hunter_render_seed_cards(seeds_result.get("seeds", []))

    if seeds_result.get("market_warning"):
        st.warning(f"⚠ **시장 위험 안내:** {seeds_result['market_warning']}")
    if seeds_result.get("recommendation"):
        st.markdown(f"**추천:** {seeds_result['recommendation']}")
    st.caption(seeds_result.get("next_step", ""))

    col_back, col_reset = st.columns(2)
    with col_back:
        _hunter_back_button("3")
    with col_reset:
        _hunter_reset_button("3")


# ────────────────────────────────────────────────────────────
# 입구 4 — What if (2턴: 가설 분석+4함정+톤3분기 → 시드)
# ────────────────────────────────────────────────────────────
def _hunter_entry_4_whatif():
    small_meta(
        "가설 분석 + 4대 함정 경고 + 톤 3분기 → 시드 후보 3개의 2턴 구조입니다. "
        "가설의 매력이 결핍/상실/세계 변형 어느 쪽 충족 판타지인지 BJND 보조 진단합니다."
    )

    stage_data = st.session_state.setdefault("hunter_stage_data", {})

    # 턴 1: 입력
    whatif_input = stage_data.get("entry4_whatif_input", "")
    if not whatif_input:
        if st.session_state.get("hunter_classified", {}).get("primary_entry", {}).get("entry_id") == 4:
            whatif_input = st.session_state.get("hunter_input", "")

        st.markdown("### 1단계 — 어떤 'What if' 가설이 떠오르셨나요?")
        with st.form("entry4_whatif_form"):
            whatif = st.text_area(
                "가설 입력",
                value=whatif_input,
                height=80,
                placeholder="예: 로또 1등 + 일주일 시간 루프 / AI가 인간을 사랑한다면 / 죽은 자가 살아 돌아온다면",
            )
            submitted = st.form_submit_button("→ 가설 분석 시작", type="primary", use_container_width=True)
        if submitted and whatif.strip():
            stage_data["entry4_whatif_input"] = whatif.strip()
            st.session_state["hunter_stage_data"] = stage_data
            st.rerun()
        _hunter_back_button("4")
        return

    st.markdown(f"**작가의 가설:** _{whatif_input}_")
    st.markdown("---")

    # 턴 2: 가설 분석 + 4함정 + 톤 3분기 + 진단 질문
    diagnosis = stage_data.get("entry4_diagnosis")
    if not diagnosis:
        if st.button("❓ 가설 분석 + 4대 함정 + 톤 3분기", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 가설 분석 + 4함정 + 톤3분기 + 진단 질문 생성 중... (30~50초)"):
                prompt_text = P.HUNTER_ENTRY_4_DIAGNOSIS_PROMPT.format(whatif_input=whatif_input)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry4_diagnosis"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("4")
        with col_reset:
            _hunter_reset_button("4")
        return

    # 가설 분석 표시
    analysis = diagnosis.get("hypothesis_analysis", {})
    st.markdown("### ❓ 가설 분석")
    st.markdown(f"**핵심 동력:** {analysis.get('core_dynamic', '')}")
    st.markdown(f"**잠정 BJND:** {analysis.get('bjnd_provisional', '')}")
    examples = analysis.get("examples", [])
    if examples:
        st.caption("유사 가설 작품: " + " / ".join(examples))

    # 4대 함정 경고
    traps = diagnosis.get("four_traps_warning", {})
    if traps:
        st.markdown("### ⚠ 4대 함정 경고")
        st.markdown(f"- **함정 1 (가설 의존성):** {traps.get('trap_1_hypothesis_crutch', '')}")
        st.markdown(f"- **함정 2 (룰 위반):** {traps.get('trap_2_rule_violation', '')}")
        st.markdown(f"- **함정 3 (결말 회피):** {traps.get('trap_3_ending_avoidance', '')}")
        st.markdown(f"- **함정 4 (일회성):** {traps.get('trap_4_one_shot', '')}")

    # 톤 3분기
    tones = diagnosis.get("three_tones", {})
    if tones:
        st.markdown("### 🎭 톤 3분기")
        col_a, col_b, col_c = st.columns(3)
        for col, key, label in [(col_a, "tone_a_comedy", "코미디"), (col_b, "tone_b_drama", "드라마/멜로"), (col_c, "tone_c_thriller", "스릴러/누아르")]:
            t = tones.get(key, {})
            with col:
                st.markdown(f"**{label}**")
                st.caption(t.get("description", ""))
                if t.get("example"):
                    st.caption(f"예: {t['example']}")

    st.markdown("---")

    # 진단 답변
    diagnosis_answers = stage_data.get("entry4_diagnosis_answers")
    if not diagnosis_answers:
        intro = "가설 분석을 보셨으니 이제 작가의 직감을 진단합니다."
        answers = _hunter_render_questions(
            diagnosis.get("diagnosis_questions", []),
            "e4_diag",
            intro_text=intro,
        )
        if st.button("→ 시드 후보 3개 빌드", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry4_diagnosis_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("4")
        with col_reset:
            _hunter_reset_button("4")
        return

    # 시드 빌드
    seeds_result = stage_data.get("entry4_seeds")
    if not seeds_result:
        if st.button("🌱 시드 후보 3개 생성 (Opus 4.7)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 가설 시드 3개 빌드 중... (40~60초)"):
                prompt_text = P.HUNTER_ENTRY_4_SEEDS_PROMPT.format(
                    whatif_input=whatif_input,
                    hypothesis_analysis=json.dumps(analysis, ensure_ascii=False, indent=2),
                    diagnosis_answers=json.dumps(diagnosis_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry4_seeds"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("4")
        with col_reset:
            _hunter_reset_button("4")
        return

    # 시드 표시
    st.markdown(f"### 🌱 시드 후보 3개 — 선택 톤: **{seeds_result.get('selected_tone', '')}**")
    bjnd_essence = seeds_result.get("bjnd_essence", {})
    if bjnd_essence:
        st.markdown(f"**BJND 본질:** {bjnd_essence.get('type', '')} — {bjnd_essence.get('explanation', '')}")
    rule_lock = seeds_result.get("hypothesis_rule_lock", "")
    if rule_lock:
        st.markdown(f'<div class="callout"><b>가설 룰 LOCK:</b> {rule_lock}</div>', unsafe_allow_html=True)
    if seeds_result.get("synthesis"):
        st.markdown(f"**본질 종합:** {seeds_result['synthesis']}")

    _hunter_render_seed_cards(seeds_result.get("seeds", []), bjnd_essence=bjnd_essence)

    if seeds_result.get("recommendation"):
        st.markdown(f"**추천:** {seeds_result['recommendation']}")
    st.caption(seeds_result.get("next_step", ""))

    col_back, col_reset = st.columns(2)
    with col_back:
        _hunter_back_button("4")
    with col_reset:
        _hunter_reset_button("4")


# ────────────────────────────────────────────────────────────
# 입구 5 — 사실 (2턴: 캔버스+5시점 → 시드)
# ────────────────────────────────────────────────────────────
def _hunter_entry_5_fact():
    small_meta(
        "사실 캔버스 + 5시점 발굴 + BJND 보조 진단 → 시드 후보 3개의 2턴 구조입니다. "
        "사실이 인물에게 만든 게 결핍/상실 어느 순간인지 진단하고, 5시점 중 어느 자리에서 들어갈지 결정합니다."
    )

    stage_data = st.session_state.setdefault("hunter_stage_data", {})

    # 턴 1: 입력
    fact_input = stage_data.get("entry5_fact_input", "")
    if not fact_input:
        if st.session_state.get("hunter_classified", {}).get("primary_entry", {}).get("entry_id") == 5:
            fact_input = st.session_state.get("hunter_input", "")

        st.markdown("### 1단계 — 어떤 역사·실화·뉴스를 작품화하고 싶으신가요?")
        with st.form("entry5_fact_form"):
            fact = st.text_area(
                "사실 입력",
                value=fact_input,
                height=80,
                placeholder="예: 1945.8.15 일본인 / 세월호 이후 / IMF 외환위기 / 5.18 광주 / n번방 사건",
            )
            submitted = st.form_submit_button("→ 사실 캔버스 펼침", type="primary", use_container_width=True)
        if submitted and fact.strip():
            stage_data["entry5_fact_input"] = fact.strip()
            st.session_state["hunter_stage_data"] = stage_data
            st.rerun()
        _hunter_back_button("5")
        return

    st.markdown(f"**작가의 사실:** _{fact_input}_")
    st.markdown("---")

    # 턴 2: 사실 캔버스 + 5시점 + 진단 질문
    diagnosis = stage_data.get("entry5_diagnosis")
    if not diagnosis:
        if st.button("📜 사실 캔버스 + 5시점 펼침", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Sonnet이 사실 캔버스 + 5시점 + 진단 질문 생성 중... (30~50초)"):
                prompt_text = P.HUNTER_ENTRY_5_DIAGNOSIS_PROMPT.format(fact_input=fact_input)
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_SONNET)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry5_diagnosis"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("5")
        with col_reset:
            _hunter_reset_button("5")
        return

    # 사실 캔버스 표시
    canvas = diagnosis.get("fact_canvas", {})
    if canvas:
        st.markdown("### 📜 사실 캔버스")
        st.markdown(f"**무슨 일:** {canvas.get('what_happened', '')}")
        st.markdown(f"**시점/공간:** {canvas.get('time_space', '')}")
        st.markdown(f"**핵심 인물 유형:** {canvas.get('key_figures', '')}")
        st.markdown(f"**구조적 원인:** {canvas.get('structural_cause', '')}")
        st.markdown(f"**결과/여파:** {canvas.get('consequences', '')}")
        less_known = canvas.get("less_known_details", [])
        if less_known:
            with st.expander("덜 알려진 디테일"):
                for d in less_known:
                    st.markdown(f"- {d}")

    # 5시점 표시
    viewpoints = diagnosis.get("five_viewpoints", {})
    if viewpoints:
        st.markdown("### 🔭 5시점")
        for key, label in [("direct", "직접"), ("pre_event", "직전"), ("post_event", "직후"), ("peripheral", "주변"), ("generational", "후세")]:
            v = viewpoints.get(key, {})
            with st.expander(f"**{label}** — {v.get('definition', '')}"):
                st.markdown(f"<span style='color:#2E7D32;'>장점: {v.get('advantage', '')}</span>", unsafe_allow_html=True)
                st.markdown(f"<span style='color:#C62828;'>위험: {v.get('risk', '')}</span>", unsafe_allow_html=True)
                if v.get("example"):
                    st.caption(f"예: {v['example']}")

    st.markdown("---")

    # 진단 답변
    diagnosis_answers = stage_data.get("entry5_diagnosis_answers")
    if not diagnosis_answers:
        intro = "사실 캔버스와 5시점을 보셨으니 이제 작가의 진입점을 진단합니다."
        answers = _hunter_render_questions(
            diagnosis.get("diagnosis_questions", []),
            "e5_diag",
            intro_text=intro,
        )
        if st.button("→ 시드 후보 3개 빌드", type="primary", use_container_width=True):
            if all(answers.values()):
                stage_data["entry5_diagnosis_answers"] = answers
                st.session_state["hunter_stage_data"] = stage_data
                st.rerun()
            else:
                st.warning("5개 질문 모두 답해주세요.")
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("5")
        with col_reset:
            _hunter_reset_button("5")
        return

    # 시드 빌드
    seeds_result = stage_data.get("entry5_seeds")
    if not seeds_result:
        if st.button("🌱 시드 후보 3개 생성 (Opus 4.7)", type="primary", use_container_width=True):
            client = get_anthropic_client()
            with st.spinner("Opus가 사실 기반 시드 3개 빌드 중... (40~60초)"):
                prompt_text = P.HUNTER_ENTRY_5_SEEDS_PROMPT.format(
                    fact_input=fact_input,
                    fact_canvas=json.dumps(canvas, ensure_ascii=False, indent=2),
                    diagnosis_answers=json.dumps(diagnosis_answers, ensure_ascii=False, indent=2),
                )
                result = call_claude(client, prompt_text, ANTHROPIC_MODEL_OPUS)
                if result.get("_parse_error"):
                    st.error("응답 파싱 실패")
                    with st.expander("Raw 응답"):
                        st.text(result.get("_raw", ""))
                else:
                    stage_data["entry5_seeds"] = result
                    st.session_state["hunter_stage_data"] = stage_data
                    st.rerun()
        col_back, col_reset = st.columns(2)
        with col_back:
            _hunter_back_button("5")
        with col_reset:
            _hunter_reset_button("5")
        return

    # 시드 표시
    st.markdown(f"### 🌱 시드 후보 3개")
    st.markdown(f"**선택 시점:** {seeds_result.get('selected_viewpoint', '')} · **선택 각도:** {seeds_result.get('selected_angle', '')}")
    bjnd_essence = seeds_result.get("bjnd_essence", {})
    if bjnd_essence:
        st.markdown(f"**BJND 본질:** {bjnd_essence.get('type', '')} — {bjnd_essence.get('explanation', '')}")
    ethics = seeds_result.get("ethics_lock", "")
    if ethics:
        st.warning(f"⚖ **윤리 LOCK:** {ethics}")
    if seeds_result.get("synthesis"):
        st.markdown(f"**본질 종합:** {seeds_result['synthesis']}")

    _hunter_render_seed_cards(seeds_result.get("seeds", []), bjnd_essence=bjnd_essence)

    if seeds_result.get("recommendation"):
        st.markdown(f"**추천:** {seeds_result['recommendation']}")
    add_research = seeds_result.get("additional_research_needed", [])
    if add_research:
        st.markdown("**추가 리서치 권장:**")
        for r in add_research:
            st.markdown(f"- {r}")
    st.caption(seeds_result.get("next_step", ""))

    col_back, col_reset = st.columns(2)
    with col_back:
        _hunter_back_button("5")
    with col_reset:
        _hunter_reset_button("5")


# ═══════════════════════════════════════════════════════════
# MAIN ROUTING (v2.0 — 3-way mode dispatch)
# ═══════════════════════════════════════════════════════════
mode = st.session_state.get("mode", "HOME")

if mode == "HOME":
    page_home()

elif mode == "HUNTER":
    entry = st.session_state.get("hunter_entry")
    if entry is None:
        page_hunter_select()
    else:
        page_hunter_entry(entry)

elif mode == "TRIAGE":
    # HUNTER에서 시드가 전달된 경우 안내 배너 노출
    if st.session_state.get("seed_loaded_from_hunter"):
        st.success(
            "🔗 HUNTER 트랙에서 발굴한 시드가 Stage 1에 자동 입력되었습니다. "
            "내용을 확인하시고 진단을 시작하세요."
        )

    render_stepper(st.session_state.get("current_stage", 1))

    stage = st.session_state.get("current_stage", 1)

    if stage == 1:
        page_stage_1()
    elif stage == 2:
        if not st.session_state.get("stage_1_input"):
            st.warning("Stage 1을 먼저 완료해주세요.")
        else:
            page_stage_2()
    elif stage == 3:
        if not st.session_state.get("stage_2_logline") or not st.session_state.get("selected_logline"):
            st.warning("Stage 2를 먼저 완료해주세요.")
        else:
            page_stage_3()
    elif stage == 4:
        if not st.session_state.get("stage_3_hook"):
            st.warning("Stage 3을 먼저 완료해주세요.")
        else:
            page_stage_4()
    elif stage == 5:
        if not st.session_state.get("stage_4_format"):
            st.warning("Stage 4를 먼저 완료해주세요.")
        else:
            page_stage_5()
    elif stage == 6:
        if not st.session_state.get("stage_5_reference"):
            st.warning("Stage 5를 먼저 완료해주세요.")
        else:
            page_stage_6()
    elif stage == 7:
        if not st.session_state.get("stage_6_market"):
            st.warning("Stage 6을 먼저 완료해주세요.")
        else:
            page_stage_7()

else:
    # 알 수 없는 모드 — 안전 fallback
    st.warning("알 수 없는 모드입니다. 홈으로 돌아갑니다.")
    st.session_state["mode"] = "HOME"
    st.rerun()

st.markdown("---")
st.caption(f"© 2026 BLUE JEANS PICTURES · Idea Engine {ENGINE_VERSION}")
